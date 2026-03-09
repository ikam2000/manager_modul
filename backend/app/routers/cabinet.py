# -*- coding: utf-8 -*-
"""Личный кабинет: сущности, документы, уведомления, настройки."""

import base64
import io
import logging
import secrets
from pathlib import Path

from fastapi import APIRouter, Depends, Query, HTTPException, UploadFile, File, Form
from fastapi.responses import Response, FileResponse
from pydantic import BaseModel, EmailStr
from sqlalchemy import select, func, or_, and_
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.database import get_db
from app.dependencies import get_current_user, get_current_admin, get_user_company_ids, get_user_company_id, require_trial_or_subscription
from app.models.user import User, UserCompany, Company, Role
from app.models.entity import (
    Category,
    SubCategory,
    Nomenclature,
    Supplier,
    Manufacturer,
    Customer,
    Supply,
    Contract,
    ContractAppendix,
)
from app.models.user_permissions import UserCompanyPermissions, PERMISSION_KEYS
from app.models.subscription import Subscription, Plan, Invoice
from app.models.saved_qr import SavedQrCode
from app.models.api_key import ApiKey
from app.models.webhook import Webhook
from app.models.support_ticket import SupportTicket, TicketReply
from app.models.suggestion import Suggestion, SuggestionUpdate
from app.models.support_attachment import TicketAttachment, SuggestionAttachment
from app.models.notification import Notification
from app.models.document import Document
from app.security import hash_password, hash_api_key
from app.services.inn_lookup import lookup_by_inn

router = APIRouter()


class CreateCompanyUserRequest(BaseModel):
    email: EmailStr
    full_name: str
    password: str


@router.get("/users")
async def list_company_users(
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Список пользователей компании (для админов) или всех (для супер-админа)."""
    if user.role == Role.super_admin:
        result = await db.execute(select(User).order_by(User.id))
        users = result.scalars().all()
    else:
        uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id))
        company_ids = {r.company_id for r in uc.scalars().all()}
        result = await db.execute(
            select(User)
            .join(UserCompany, User.id == UserCompany.user_id)
            .where(UserCompany.company_id.in_(company_ids))
        )
        users = result.unique().scalars().all()
    perm_result = await db.execute(select(UserCompanyPermissions))
    perm_map = {f"{p.user_id}_{p.company_id}": p.permissions for p in perm_result.scalars().all()}
    items = []
    for u in users:
        uc_list = (await db.execute(select(UserCompany).where(UserCompany.user_id == u.id))).scalars().all()
        perms = {"can_delete_entities": False, "can_delete_documents": False}
        if u.role in (Role.admin, Role.super_admin):
            perms = {"can_delete_entities": True, "can_delete_documents": True}
        else:
            for uc_row in uc_list:
                key = f"{u.id}_{uc_row.company_id}"
                if key in perm_map and perm_map[key]:
                    perms["can_delete_entities"] = perms["can_delete_entities"] or bool(perm_map[key].get("can_delete_entities"))
                    perms["can_delete_documents"] = perms["can_delete_documents"] or bool(perm_map[key].get("can_delete_documents"))
        items.append({
            "id": u.id,
            "email": u.email,
            "full_name": u.full_name,
            "role": u.role.value,
            "is_active": u.is_active,
            "permissions": perms,
        })
    return {"items": items}


@router.get("/users/{user_id}")
async def get_company_user(
    user_id: int,
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Получить одного пользователя с правами."""
    if user.role != Role.super_admin:
        uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id))
        my_companies = {r.company_id for r in uc.scalars().all()}
        target_uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user_id))
        target_companies = {r.company_id for r in target_uc.scalars().all()}
        if not (my_companies & target_companies):
            raise HTTPException(403, "Пользователь из другой компании")
    result = await db.execute(select(User).where(User.id == user_id))
    u = result.scalar_one_or_none()
    if not u:
        raise HTTPException(404, "Пользователь не найден")
    perms = {"can_delete_entities": False, "can_delete_documents": False}
    if u.role in (Role.admin, Role.super_admin):
        perms = {"can_delete_entities": True, "can_delete_documents": True}
    else:
        perm_records = (
            await db.execute(
                select(UserCompanyPermissions).where(UserCompanyPermissions.user_id == user_id)
            )
        ).scalars().all()
        for pr in perm_records:
            if pr.permissions:
                perms["can_delete_entities"] = perms["can_delete_entities"] or bool(pr.permissions.get("can_delete_entities"))
                perms["can_delete_documents"] = perms["can_delete_documents"] or bool(pr.permissions.get("can_delete_documents"))
    return {
        "id": u.id,
        "email": u.email,
        "full_name": u.full_name,
        "role": u.role.value,
        "is_active": u.is_active,
        "permissions": perms,
    }


@router.post("/users")
async def create_company_user(
    data: CreateCompanyUserRequest,
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Создать пользователя в компании (для админов)."""
    if len(data.password) < 8:
        raise HTTPException(400, "Пароль минимум 8 символов")
    existing = await db.execute(select(User).where(User.email == data.email))
    if existing.scalar_one_or_none():
        raise HTTPException(400, "Email уже занят")
    uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id).limit(1))
    row = uc.scalar_one_or_none()
    if not row:
        raise HTTPException(403, "Нет компании")
    company_id = row.company_id
    new_user = User(
        email=data.email,
        full_name=data.full_name,
        hashed_password=hash_password(data.password),
        role=Role.user,
    )
    db.add(new_user)
    await db.flush()
    db.add(UserCompany(user_id=new_user.id, company_id=company_id, role=Role.user))
    await db.commit()
    return {"id": new_user.id, "email": new_user.email}


class UpdateCompanyUserRequest(BaseModel):
    is_active: bool | None = None
    role: str | None = None


@router.patch("/users/{user_id}")
async def update_company_user(
    user_id: int,
    data: UpdateCompanyUserRequest,
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Обновить пользователя (включить/отключить, роль — только супер-админ)."""
    if user.role != Role.super_admin:
        uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id))
        my_companies = {r.company_id for r in uc.scalars().all()}
        target_uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user_id))
        target_companies = {r.company_id for r in target_uc.scalars().all()}
        if not (my_companies & target_companies):
            raise HTTPException(403, "Пользователь из другой компании")
    result = await db.execute(select(User).where(User.id == user_id))
    target = result.scalar_one_or_none()
    if not target:
        raise HTTPException(404, "Пользователь не найден")
    if target.role == Role.super_admin:
        raise HTTPException(403, "Нельзя изменить супер-админа")
    if data.is_active is not None:
        target.is_active = data.is_active
    if data.role is not None and user.role == Role.super_admin and data.role in ("user", "admin", "trader"):
        target.role = Role(data.role)
    await db.commit()
    return {"ok": True}


@router.delete("/users/{user_id}")
async def delete_company_user(
    user_id: int,
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Удалить пользователя."""
    if user.role != Role.super_admin:
        uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id))
        my_companies = {r.company_id for r in uc.scalars().all()}
        target_uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user_id))
        target_companies = {r.company_id for r in target_uc.scalars().all()}
        if not (my_companies & target_companies):
            raise HTTPException(403, "Пользователь из другой компании")
    result = await db.execute(select(User).where(User.id == user_id))
    target = result.scalar_one_or_none()
    if not target:
        raise HTTPException(404, "Пользователь не найден")
    if target.role == Role.super_admin:
        raise HTTPException(403, "Нельзя удалить супер-админа")
    await db.delete(target)
    await db.commit()
    return {"ok": True}


class UpdatePermissionsRequest(BaseModel):
    can_delete_entities: bool | None = None
    can_delete_documents: bool | None = None


@router.patch("/users/{user_id}/permissions")
async def update_user_permissions(
    user_id: int,
    data: UpdatePermissionsRequest,
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Обновить права пользователя (удаление номенклатуры, документов)."""
    if user.role != Role.super_admin:
        uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id))
        my_companies = {r.company_id for r in uc.scalars().all()}
        target_uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user_id))
        target_rows = target_uc.scalars().all()
        target_companies = {r.company_id for r in target_rows}
        if not (my_companies & target_companies):
            raise HTTPException(403, "Пользователь из другой компании")
        company_ids = list(my_companies & target_companies)
    else:
        target_uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user_id))
        target_rows = target_uc.scalars().all()
        company_ids = [r.company_id for r in target_rows]
    if not company_ids:
        raise HTTPException(404, "Пользователь не привязан к компании")
    target_user = await db.execute(select(User).where(User.id == user_id))
    if not target_user.scalar_one_or_none():
        raise HTTPException(404, "Пользователь не найден")
    for cid in company_ids:
        r = await db.execute(
            select(UserCompanyPermissions).where(
                UserCompanyPermissions.user_id == user_id,
                UserCompanyPermissions.company_id == cid,
            )
        )
        rec = r.scalar_one_or_none()
        if not rec:
            rec = UserCompanyPermissions(user_id=user_id, company_id=cid, permissions={})
            db.add(rec)
            await db.flush()
        perms = dict(rec.permissions or {})
        if data.can_delete_entities is not None:
            perms["can_delete_entities"] = data.can_delete_entities
        if data.can_delete_documents is not None:
            perms["can_delete_documents"] = data.can_delete_documents
        rec.permissions = {k: v for k, v in perms.items() if k in PERMISSION_KEYS}
    await db.commit()
    return {"ok": True}


def _storage_root() -> Path:
    p = get_settings().storage_path
    if not p.is_absolute():
        p = Path(__file__).resolve().parents[1] / "storage"
    p.mkdir(parents=True, exist_ok=True)
    return p


REQUISITES_FIELDS = ["name", "inn", "kpp", "ogrn", "legal_address", "address", "phone", "email", "contact_person", "bank_name", "bank_bik", "bank_account", "bank_corr", "payment_purpose"]


@router.get("/inn-lookup")
async def inn_lookup(
    inn: str = Query(..., min_length=10),
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Поиск организации по ИНН (entity_registry + DaData). Данные для подстановки в формы."""
    data = await lookup_by_inn(inn, db)
    if not data:
        return {"found": False, "data": None}
    return {"found": True, "data": data}


@router.get("/company")
async def get_company(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Реквизиты компании пользователя."""
    uc = await db.execute(select(UserCompany, Company).join(Company, UserCompany.company_id == Company.id).where(UserCompany.user_id == user.id).limit(1))
    row = uc.first()
    if not row:
        return {"company": None}
    _, company = row
    return {
        "company": {
            "id": company.id,
            "name": company.name,
            "logo_url": "/api/cabinet/company/logo" if getattr(company, "logo_url", None) else None,
            "inn": company.inn,
            "kpp": company.kpp,
            "ogrn": getattr(company, "ogrn", None),
            "legal_address": company.legal_address,
            "address": getattr(company, "address", None),
            "phone": getattr(company, "phone", None),
            "email": getattr(company, "email", None),
            "contact_person": getattr(company, "contact_person", None),
            "bank_name": getattr(company, "bank_name", None),
            "bank_bik": getattr(company, "bank_bik", None),
            "bank_account": getattr(company, "bank_account", None),
            "bank_corr": getattr(company, "bank_corr", None),
            "payment_purpose": getattr(company, "payment_purpose", None),
        }
    }


class CompanyRequisitesUpdate(BaseModel):
    name: str | None = None
    inn: str | None = None
    kpp: str | None = None
    ogrn: str | None = None
    legal_address: str | None = None
    address: str | None = None
    phone: str | None = None
    email: str | None = None
    contact_person: str | None = None
    bank_name: str | None = None
    bank_bik: str | None = None
    bank_account: str | None = None
    bank_corr: str | None = None
    payment_purpose: str | None = None


@router.patch("/company")
async def update_company(
    data: CompanyRequisitesUpdate,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Обновить реквизиты компании."""
    uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id).limit(1))
    row = uc.scalar_one_or_none()
    if not row:
        raise HTTPException(403, "Нет компании")
    company = await db.get(Company, row.company_id)
    if not company:
        raise HTTPException(404, "Компания не найдена")
    for k, v in data.model_dump(exclude_unset=True).items():
        if hasattr(company, k):
            setattr(company, k, v)
    await db.commit()
    return {"ok": True}


@router.post("/avatar")
async def upload_avatar(
    file: UploadFile = File(...),
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Загрузить аватар пользователя (jpg, png, webp)."""
    if not file.content_type or not file.content_type.startswith("image/"):
        raise HTTPException(400, "Только изображения (jpg, png, webp)")
    ext = Path(file.filename or "").suffix.lower() or ".png"
    if ext not in (".jpg", ".jpeg", ".png", ".webp"):
        ext = ".png"
    content = await file.read()
    if len(content) > 2 * 1024 * 1024:
        raise HTTPException(400, "Размер до 2 МБ")
    root = _storage_root()
    rel = f"avatars/user_{user.id}{ext}"
    path = root / rel
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "wb") as f:
        f.write(content)
    user.avatar_url = rel
    await db.commit()
    b64 = base64.b64encode(content).decode("ascii")
    mime = "image/png" if ext == ".png" else "image/jpeg" if ext in (".jpg", ".jpeg") else "image/webp"
    return {"avatar_url": "/api/cabinet/avatar", "avatar_data": f"data:{mime};base64,{b64}"}


@router.get("/avatar")
async def get_avatar(
    user: User = Depends(require_trial_or_subscription),
):
    """Получить аватар текущего пользователя."""
    if not getattr(user, "avatar_url", None):
        raise HTTPException(404, "Аватар не загружен")
    root = _storage_root()
    path = root / user.avatar_url
    if not path.exists():
        raise HTTPException(404, "Файл аватара не найден")
    ext = path.suffix.lower()
    media = "image/png" if ext == ".png" else "image/jpeg" if ext in (".jpg", ".jpeg") else "image/webp" if ext == ".webp" else "image/png"
    return FileResponse(path, media_type=media, headers={"Cache-Control": "no-cache, no-store"})


@router.post("/company/logo")
async def upload_logo(
    file: UploadFile = File(...),
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Загрузить логотип компании (png, svg — прозрачный или белый фон)."""
    ct = file.content_type or ""
    ext = Path(file.filename or "").suffix.lower() or ".png"
    if ext not in (".png", ".svg", ".jpg", ".jpeg", ".webp") and "image/" not in ct:
        raise HTTPException(400, "Только изображения (png, svg, jpg, webp)")
    content = await file.read()
    if len(content) > 2 * 1024 * 1024:
        raise HTTPException(400, "Размер до 2 МБ")
    uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id).limit(1))
    row = uc.scalar_one_or_none()
    if not row:
        raise HTTPException(403, "Нет компании")
    root = _storage_root()
    rel = f"logos/company_{row.company_id}{ext}"
    path = root / rel
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "wb") as f:
        f.write(content)
    company = await db.get(Company, row.company_id)
    if company:
        company.logo_url = rel
        await db.commit()
    return {"logo_url": "/api/cabinet/company/logo"}


class ApplyRequisitesRequest(BaseModel):
    mapping: dict[str, str]  # field -> value from recognized


@router.post("/company/recognize-requisites")
async def recognize_requisites(
    file: UploadFile = File(...),
    user: User = Depends(require_trial_or_subscription),
):
    """Распознать реквизиты из файла (PDF, Word, картинка). Возвращает данные для сопоставления полей."""
    import re
    import io
    from pypdf import PdfReader

    content = await file.read()
    filename = file.filename or "doc"
    ext = Path(filename).suffix.lower()
    text = ""
    detected: dict[str, str] = {}

    try:
        if ext in (".pdf",) or (file.content_type or "").endswith("pdf"):
            reader = PdfReader(io.BytesIO(content))
            for p in reader.pages:
                t = p.extract_text()
                if t:
                    text += t + "\n"
        elif ext in (".docx",) or "wordprocessingml" in (file.content_type or ""):
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            for p in doc.paragraphs:
                if p.text.strip():
                    text += p.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    text += " ".join(c.text for c in row.cells) + "\n"
        elif ext in (".jpg", ".jpeg", ".png"):
            try:
                import pytesseract
                from PIL import Image
                img = Image.open(io.BytesIO(content))
                text = pytesseract.image_to_string(img, lang="rus+eng") or ""
            except ImportError:
                text = ""
    except Exception:
        text = ""

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    raw = " ".join(lines).lower()

    patterns = [
        (r"инн[\s:]*(\d{10,12})", "inn"),
        (r"кпп[\s:]*(\d{9})", "kpp"),
        (r"огрн[\s:]*(\d{13,15})", "ogrn"),
        (r"расч[\.\s]*счет[\s:]*(\d{20})", "bank_account"),
        (r"р/с[\s:]*(\d{20})", "bank_account"),
        (r"счет[\s:]*№?\s*(\d{20})", "bank_account"),
        (r"бик[\s:]*(\d{9})", "bank_bik"),
        (r"банк[\s:]+([^\n]+?)(?=\s+бик|\s+инн|\s+кпп|$)", "bank_name"),
        (r"корр[\.\s]*счет[\s:]*(\d{20})", "bank_corr"),
        (r"к/с[\s:]*(\d{20})", "bank_corr"),
        (r"адрес[\s:]+([^\n]+)", "legal_address"),
        (r"юридический адрес[\s:]+([^\n]+)", "legal_address"),
        (r"наименование[\s:]+([^\n]+)", "name"),
        (r"полное наименование[\s:]+([^\n]+)", "name"),
    ]
    for pat, key in patterns:
        m = re.search(pat, raw, re.I | re.DOTALL)
        if m and key not in detected:
            val = m.group(1).strip()
            if len(val) < 500:
                detected[key] = val[:255] if key in ("legal_address", "bank_name", "name", "payment_purpose") else val[:50]

    detected_columns = [{"index": i, "name": k, "example": v} for i, (k, v) in enumerate(detected.items())]
    return {"detected": detected, "detected_columns": detected_columns, "raw_preview": text[:1000]}


@router.post("/company/apply-requisites")
async def apply_requisites(
    data: ApplyRequisitesRequest,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Применить сопоставленные реквизиты к компании."""
    uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id).limit(1))
    row = uc.scalar_one_or_none()
    if not row:
        raise HTTPException(403, "Нет компании")
    company = await db.get(Company, row.company_id)
    if not company:
        raise HTTPException(404, "Компания не найдена")
    for k, v in data.mapping.items():
        if k in REQUISITES_FIELDS and v is not None:
            setattr(company, k, str(v).strip()[:512] if len(str(v)) > 100 else str(v).strip()[:100])
    await db.commit()
    return {"ok": True}


@router.get("/company/logo")
async def get_logo(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Получить логотип компании."""
    uc = await db.execute(select(UserCompany, Company).join(Company, UserCompany.company_id == Company.id).where(UserCompany.user_id == user.id).limit(1))
    row = uc.first()
    if not row or not getattr(row[1], "logo_url", None):
        raise HTTPException(404, "Логотип не загружен")
    _, company = row
    root = _storage_root()
    path = root / company.logo_url
    if not path.exists():
        raise HTTPException(404, "Файл логотипа не найден")
    return FileResponse(path, media_type="image/png")


@router.get("/entities")
async def list_entities(
    user: User = Depends(require_trial_or_subscription),
    entity_type: str | None = Query(None),
    limit: int = Query(50, le=200),
    offset: int = Query(0, ge=0),
):
    """Список сущностей с фильтрами."""
    return {"items": [], "total": 0}


def _generate_supply_template_xlsx() -> bytes:
    """Шаблон XLSX для отгрузочной документации."""
    from openpyxl import Workbook
    from openpyxl.styles import Font
    wb = Workbook()
    ws = wb.active
    ws.title = "Ведомость"
    ws.append(["Ведомость грузового места / Packing note", "", "", "", "", ""])
    ws.append(["Договор / Contract", "", "", "", "", ""])
    ws.append(["Продавец / Seller", "", "", "", "Покупатель / Buyer", ""])
    ws.append(["Адрес отгрузки / Address of pickup", "", "", "", "Адрес поставки / Delivery address", ""])
    ws.append(["Дата отгрузки", "", "", "", "", ""])
    ws.append([])
    headers = ["№", "Таговый номер / Tag number", "Артикул / Vendor code", "Наименование товара / Product name", "Кол-во / Qty", "Производитель / Manufacturer", "№ груз. места"]
    ws.append(headers)
    ws.append([1, "ТН-001", "МП-100", "Манометр показывающий 0-1 МПа", 2, "Завод Манометр", "GM-001"])
    ws.append([2, "ТН-002", "МП-160", "Манометр показывающий 0-2.5 МПа", 1, "Завод Манометр", "GM-001"])
    ws.append([3, "", "", "", "", "", ""])
    for col in range(1, len(headers) + 1):
        ws.cell(1, col).font = Font(bold=True)
        ws.cell(9, col).font = Font(bold=True)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


@router.get("/template/supplies")
async def download_supply_template(
    user: User = Depends(require_trial_or_subscription),
):
    """Скачать шаблон XLSX для отгрузочной документации (альтернативный маршрут)."""
    content = _generate_supply_template_xlsx()
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="shablon_otgruzochnaya_vedomost.xlsx"'},
    )


logger = logging.getLogger(__name__)


@router.get("/entity-tree")
async def get_entity_tree(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
    company_id: int | None = Query(None, description="Для супер-админа: фильтр по компании"),
):
    """Дерево сущностей для раздела документов: категории с подкатегориями и номенклатурой,
    поставщики, производители, поставки, договоры. Фильтр по компании пользователя."""
    try:
        return await _get_entity_tree_impl(user, db, company_id)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("entity-tree 500: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


async def _get_entity_tree_impl(user, db, company_id):
    company_ids = await get_user_company_ids(user, db)
    if company_ids is not None and not company_ids:
        return {
            "categories": [],
            "suppliers": [],
            "manufacturers": [],
            "customers": [],
            "supplies": [],
            "contracts": [],
        }
    filter_ids = [company_id] if company_id is not None and user.role == Role.super_admin else company_ids

    # Categories with subcategories and nomenclature
    cat_q = select(Category).order_by(Category.name)
    if filter_ids is not None:
        cat_q = cat_q.where(Category.company_id.in_(filter_ids))
    cats_result = await db.execute(cat_q)
    categories_rows = cats_result.scalars().all()

    subcat_q = select(SubCategory).join(Category, SubCategory.category_id == Category.id)
    if filter_ids is not None:
        subcat_q = subcat_q.where(Category.company_id.in_(filter_ids))
    subcat_q = subcat_q.order_by(SubCategory.name)
    subcats_result = await db.execute(subcat_q)
    subcats_all = subcats_result.scalars().all()

    nom_q = select(Nomenclature).where(Nomenclature.is_deleted == False)
    if filter_ids is not None:
        nom_q = nom_q.where(Nomenclature.company_id.in_(filter_ids))
    nom_result = await db.execute(nom_q)
    noms_all = nom_result.scalars().all()

    supply_q = select(Supply).where(Supply.is_deleted == False).order_by(Supply.created_at.desc())
    if filter_ids is not None:
        supply_q = supply_q.where(Supply.company_id.in_(filter_ids))
    supplies = (await db.execute(supply_q)).scalars().all()

    subcat_by_cat: dict[int, list] = {c.id: [] for c in categories_rows}
    for sc in subcats_all:
        if sc.category_id in subcat_by_cat:
            subcat_by_cat[sc.category_id].append(sc)

    nom_by_subcat: dict[int, list] = {}
    nom_without_subcat_by_cat: dict[int, list] = {c.id: [] for c in categories_rows}
    for n in noms_all:
        if n.subcategory_id:
            nom_by_subcat.setdefault(n.subcategory_id, []).append(n)
        elif n.category_id and n.category_id in nom_without_subcat_by_cat:
            nom_without_subcat_by_cat[n.category_id].append(n)
    for subcat_list in subcat_by_cat.values():
        for sc in subcat_list:
            sc._nomenclature = nom_by_subcat.get(sc.id, [])

    categories = []
    for c in categories_rows:
        subcats_data = []
        for sc in subcat_by_cat.get(c.id, []):
            noms = [
                {"id": n.id, "code": n.code or "", "name": n.name}
                for n in getattr(sc, "_nomenclature", [])
            ]
            subcats_data.append({"id": sc.id, "name": sc.name, "nomenclature": noms})
        other_noms = nom_without_subcat_by_cat.get(c.id, [])
        if other_noms:
            subcats_data.append({
                "id": -c.id,
                "name": "Прочее",
                "nomenclature": [{"id": n.id, "code": n.code or "", "name": n.name} for n in other_noms],
            })
        categories.append({"id": c.id, "name": c.name, "subcategories": subcats_data})
    # Добавим номенклатуру без подкатегории в отдельную группу? В ТЗ: "each subcat with nomenclature items"
    # Оставляем только в подкатегориях. Номенклатура без subcategory_id не попадёт — это нормально.

    # Suppliers с количеством номенклатуры (через Supply)
    sup_q = select(Supplier).where(Supplier.is_deleted == False).order_by(Supplier.name)
    if filter_ids is not None:
        sup_q = sup_q.where(Supplier.company_id.in_(filter_ids))
    suppliers = (await db.execute(sup_q)).scalars().all()
    nom_ids_set = {n.id for n in noms_all}
    supplier_nom_ids: dict[int, set[int]] = {}
    for sp in supplies:
        if sp.supplier_id and sp.nomenclature_id and sp.nomenclature_id in nom_ids_set:
            supplier_nom_ids.setdefault(sp.supplier_id, set()).add(sp.nomenclature_id)
    suppliers_data = [
        {"id": s.id, "name": s.name, "code": s.inn or "", "count": len(supplier_nom_ids.get(s.id, set()))}
        for s in suppliers
    ]

    # Manufacturers — с fallback при ошибке схемы
    try:
        mfr_q = select(Manufacturer).where(Manufacturer.is_deleted == False).order_by(Manufacturer.name)
        if filter_ids is not None:
            mfr_q = mfr_q.where(Manufacturer.company_id.in_(filter_ids))
        manufacturers = (await db.execute(mfr_q)).scalars().all()
        manufacturers_data = [{"id": m.id, "name": m.name, "code": ""} for m in manufacturers]
    except Exception as e:
        logger.warning("entity-tree manufacturers fallback: %s", e)
        manufacturers_data = []

    # Customers (заказчики) — с fallback при ошибке схемы
    try:
        cust_q = select(Customer).where(Customer.is_deleted == False).order_by(Customer.name)
        if filter_ids is not None:
            cust_q = cust_q.where(Customer.company_id.in_(filter_ids))
        customers = (await db.execute(cust_q)).scalars().all()
        customers_data = [{"id": c.id, "name": c.name, "code": getattr(c, "inn", None) or ""} for c in customers]
    except Exception as e:
        logger.warning("entity-tree customers fallback: %s", e)
        customers_data = []

    # Supplies (flat) — supplies уже загружены выше
    supplies_data = []
    supplier_map = {s.id: s for s in suppliers}
    nom_map = {n.id: n for n in noms_all}
    for sp in supplies:
        parts = []
        if sp.supplier_id and sp.supplier_id in supplier_map:
            parts.append(supplier_map[sp.supplier_id].name)
        if sp.nomenclature_id and sp.nomenclature_id in nom_map:
            parts.append(nom_map[sp.nomenclature_id].name or nom_map[sp.nomenclature_id].code or "")
        name = " — ".join(parts) if parts else f"Поставка #{sp.id}"
        supplies_data.append({"id": sp.id, "name": name, "code": ""})

    # Contracts (flat)
    contract_q = select(Contract).where(Contract.is_deleted == False).order_by(Contract.created_at.desc())
    if filter_ids is not None:
        contract_q = contract_q.where(Contract.company_id.in_(filter_ids))
    contracts = (await db.execute(contract_q)).scalars().all()
    contracts_data = [
        {"id": c.id, "name": c.number or f"Договор #{c.id}", "code": c.number or "", "supplier_id": c.supplier_id}
        for c in contracts
    ]

    # Кастомные поля номенклатуры (уникальные ключи из extra_fields) для выпадающего списка при распознавании
    nom_custom_fields: set[str] = set()
    for n in noms_all:
        if n.extra_fields and isinstance(n.extra_fields, dict):
            nom_custom_fields.update(k for k in n.extra_fields.keys() if k and isinstance(k, str))
    nomenclature_custom_fields = sorted(nom_custom_fields)

    return {
        "categories": categories,
        "suppliers": suppliers_data,
        "manufacturers": manufacturers_data,
        "customers": customers_data,
        "supplies": supplies_data,
        "contracts": contracts_data,
        "nomenclature_custom_fields": nomenclature_custom_fields,
    }


@router.get("/documents-counts")
async def get_documents_counts(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
    company_id: int | None = Query(None, description="Для супер-админа: фильтр по компании"),
):
    """Количество документов по категориям, поставщикам, производителям для sidebar в разделе Документы."""
    company_ids = await get_user_company_ids(user, db)
    if company_ids is not None and not company_ids:
        return {"total": 0, "by_category": {}, "by_supplier": {}, "by_manufacturer": {}, "by_supply": {}}
    filter_ids = [company_id] if company_id is not None and user.role == Role.super_admin else company_ids

    # total
    total_r = await db.execute(select(func.count()).select_from(Document).where(Document.company_id.in_(filter_ids)))
    total = total_r.scalar() or 0

    # by category: документы номенклатуры в категории
    cat_q = select(Category.id).where(Category.company_id.in_(filter_ids))
    cats = [r[0] for r in (await db.execute(cat_q)).all()]
    by_category = {}
    for cid in cats:
        subcat_q = select(SubCategory.id).where(SubCategory.category_id == cid)
        subcats = [r[0] for r in (await db.execute(subcat_q)).all()]
        cat_filter = Nomenclature.category_id == cid
        subcat_filter = Nomenclature.subcategory_id.in_(subcats) if subcats else (Nomenclature.id < 0)
        nom_ids_q = select(Nomenclature.id).where(
            or_(cat_filter, subcat_filter),
            Nomenclature.is_deleted == False,
        ).where(Nomenclature.company_id.in_(filter_ids))
        nom_ids = [r[0] for r in (await db.execute(nom_ids_q)).all()]
        if nom_ids:
            cnt_q = select(func.count()).select_from(Document).where(
                Document.company_id.in_(filter_ids),
                Document.entity_type == "nomenclature",
                Document.entity_id.in_(nom_ids),
            )
            by_category[cid] = (await db.execute(cnt_q)).scalar() or 0
        else:
            by_category[cid] = 0

    # by supplier
    sup_q = select(Supplier.id).where(Supplier.company_id.in_(filter_ids))
    sups = [r[0] for r in (await db.execute(sup_q)).all()]
    by_supplier = {}
    for sid in sups:
        supply_ids_q = select(Supply.id).where(Supply.supplier_id == sid, Supply.is_deleted == False).where(Supply.company_id.in_(filter_ids))
        supply_ids = [r[0] for r in (await db.execute(supply_ids_q)).all()]
        contract_ids_q = select(Contract.id).where(Contract.supplier_id == sid, Contract.is_deleted == False).where(Contract.company_id.in_(filter_ids))
        contract_ids = [r[0] for r in (await db.execute(contract_ids_q)).all()]
        nom_ids_q = select(Supply.nomenclature_id).where(
            Supply.supplier_id == sid,
            Supply.is_deleted == False,
            Supply.nomenclature_id.isnot(None),
        ).where(Supply.company_id.in_(filter_ids))
        nom_ids = list({r[0] for r in (await db.execute(nom_ids_q)).all() if r[0]})
        conds = [
            and_(Document.entity_type == "supplier", Document.entity_id == sid),
        ]
        if supply_ids:
            conds.append(and_(Document.entity_type == "supply", Document.entity_id.in_(supply_ids)))
        if contract_ids:
            conds.append(and_(Document.entity_type == "contract", Document.entity_id.in_(contract_ids)))
        if nom_ids:
            conds.append(and_(Document.entity_type == "nomenclature", Document.entity_id.in_(nom_ids)))
        cnt_q = select(func.count()).select_from(Document).where(
            Document.company_id.in_(filter_ids),
            or_(*conds),
        )
        by_supplier[sid] = (await db.execute(cnt_q)).scalar() or 0

    # by manufacturer
    mfr_q = select(Manufacturer.id).where(Manufacturer.company_id.in_(filter_ids))
    mfrs = [r[0] for r in (await db.execute(mfr_q)).all()]
    by_manufacturer = {}
    for mid in mfrs:
        nom_ids_q = select(Nomenclature.id).where(
            Nomenclature.manufacturer_id == mid,
            Nomenclature.is_deleted == False,
        ).where(Nomenclature.company_id.in_(filter_ids))
        nom_ids = [r[0] for r in (await db.execute(nom_ids_q)).all()]
        conds = [and_(Document.entity_type == "manufacturer", Document.entity_id == mid)]
        if nom_ids:
            conds.append(and_(Document.entity_type == "nomenclature", Document.entity_id.in_(nom_ids)))
        cnt_q = select(func.count()).select_from(Document).where(
            Document.company_id.in_(filter_ids),
            or_(*conds),
        )
        by_manufacturer[mid] = (await db.execute(cnt_q)).scalar() or 0

    # by supply: документы, привязанные к поставкам
    supply_ids_q = select(Supply.id).where(Supply.is_deleted == False)
    if filter_ids is not None:
        supply_ids_q = supply_ids_q.where(Supply.company_id.in_(filter_ids))
    supply_ids = [r[0] for r in (await db.execute(supply_ids_q)).all()]
    by_supply = {}
    for sid in supply_ids:
        cnt_q = select(func.count()).select_from(Document).where(
            Document.company_id.in_(filter_ids),
            Document.entity_type == "supply",
            Document.entity_id == sid,
        )
        cnt = (await db.execute(cnt_q)).scalar() or 0
        if cnt:
            by_supply[sid] = cnt

    return {
        "total": total,
        "by_category": by_category,
        "by_supplier": by_supplier,
        "by_manufacturer": by_manufacturer,
        "by_supply": by_supply,
    }


@router.get("/label-data")
async def get_label_data(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
    nomenclature_ids: str = Query(..., description="ID номенклатуры через запятую"),
):
    """Данные для этикеток: номенклатура + поставщик, договор, спецификация, дата изготовления из последней поставки."""
    company_ids = await get_user_company_ids(user, db)
    if company_ids is not None and not company_ids:
        return {"items": []}

    ids = [int(x.strip()) for x in nomenclature_ids.split(",") if x.strip()]
    if not ids:
        return {"items": []}

    nom_q = (
        select(Nomenclature)
        .where(Nomenclature.id.in_(ids), Nomenclature.is_deleted == False)
        .where(Nomenclature.company_id.in_(company_ids))
    )
    noms_r = await db.execute(nom_q)
    noms = {n.id: n for n in noms_r.scalars().all()}

    # Последняя поставка по каждой номенклатуре
    supply_q = (
        select(Supply)
        .where(Supply.nomenclature_id.in_(ids), Supply.is_deleted == False)
        .where(Supply.company_id.in_(company_ids))
        .order_by(Supply.created_at.desc())
    )
    supplies_r = await db.execute(supply_q)
    supplies_all = supplies_r.scalars().all()

    # Одна поставка на номенклатуру (последняя)
    supply_by_nom: dict[int, Supply] = {}
    for s in supplies_all:
        if s.nomenclature_id and s.nomenclature_id not in supply_by_nom:
            supply_by_nom[s.nomenclature_id] = s

    supplier_ids = [s.supplier_id for s in supply_by_nom.values() if s.supplier_id]
    suppliers_r = await db.execute(
        select(Supplier).where(Supplier.id.in_(supplier_ids))
    )
    suppliers_map = {s.id: s for s in suppliers_r.scalars().all()}

    contract_q = select(Contract).where(
        Contract.supplier_id.in_(supplier_ids),
        Contract.is_deleted == False,
    )
    if company_ids:
        contract_q = contract_q.where(Contract.company_id.in_(company_ids))
    contracts_r = await db.execute(contract_q.order_by(Contract.date_start.desc()))
    contracts_all = contracts_r.scalars().all()
    contract_by_supplier: dict[int, Contract] = {}
    for c in contracts_all:
        if c.supplier_id and c.supplier_id not in contract_by_supplier:
            contract_by_supplier[c.supplier_id] = c

    contract_ids = [c.id for c in contract_by_supplier.values()]
    appendices_r = await db.execute(
        select(ContractAppendix)
        .where(ContractAppendix.contract_id.in_(contract_ids))
        .where(ContractAppendix.is_deleted == False)
    )
    appendices_all = appendices_r.scalars().all()
    appendix_by_contract: dict[int, ContractAppendix] = {}
    for a in appendices_all:
        if a.contract_id not in appendix_by_contract:
            appendix_by_contract[a.contract_id] = a

    items = []
    for nom_id in ids:
        nom = noms.get(nom_id)
        if not nom:
            continue
        sup = supply_by_nom.get(nom_id)
        supplier = suppliers_map.get(sup.supplier_id) if sup and sup.supplier_id else None
        contract = contract_by_supplier.get(supplier.id) if supplier else None
        appendix = appendix_by_contract.get(contract.id) if contract else None

        def _date(d) -> str | None:
            if d is None:
                return None
            return d.strftime("%d.%m.%Y") if hasattr(d, "strftime") else str(d)

        items.append({
            "nomenclature_id": nom.id,
            "product_name": nom.name or "",
            "product_code": nom.code or "",
            "object_code": nom.tag_number or nom.package_number or "",
            "supplier_name": supplier.name if supplier else "",
            "contact_phone": supplier.phone if supplier else "",
            "contract_number": contract.number if contract else "",
            "contract_date": _date(contract.date_start) if contract else None,
            "specification_number": appendix.name if appendix else "",
            "specification_date": _date(appendix.created_at) if appendix and hasattr(appendix, "created_at") else None,
            "production_date": _date(sup.production_date) if sup else None,
        })

    return {"items": items}


@router.get("/notifications")
async def list_notifications(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
    unread_only: bool = Query(False),
):
    """Уведомления пользователя."""
    q = select(Notification).where(Notification.user_id == user.id)
    if unread_only:
        q = q.where(Notification.is_read == False)
    q = q.order_by(Notification.created_at.desc()).limit(100)
    r = await db.execute(q)
    items = r.scalars().all()
    return {
        "items": [
            {
                "id": n.id,
                "title": n.title,
                "body": n.body,
                "type": n.type,
                "link": n.link,
                "is_read": n.is_read,
                "created_at": n.created_at.isoformat(),
            }
            for n in items
        ]
    }


@router.patch("/notifications/{nid}/read")
async def mark_notification_read(
    nid: int,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Отметить уведомление как прочитанное."""
    r = await db.execute(select(Notification).where(Notification.id == nid, Notification.user_id == user.id))
    n = r.scalar_one_or_none()
    if not n:
        raise HTTPException(404, "Не найдено")
    n.is_read = True
    await db.commit()
    return {"ok": True}


@router.get("/subscription")
async def get_subscription(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Текущая подписка компании пользователя."""
    uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id).limit(1))
    row = uc.scalar_one_or_none()
    if not row:
        return {"status": "trial", "plan_name": "Пробный", "expires_at": None}
    result = await db.execute(
        select(Subscription, Plan)
        .join(Plan, Subscription.plan_id == Plan.id)
        .where(Subscription.company_id == row.company_id)
        .order_by(Subscription.created_at.desc())
        .limit(1)
    )
    sub_row = result.first()
    if not sub_row:
        return {"status": "trial", "plan_name": "Пробный", "expires_at": None}
    sub, plan = sub_row
    company_id = row.company_id
    is_demo = plan.name == "Демо"
    has_limits = any(
        getattr(plan, k, None) is not None
        for k in ("max_nomenclature", "max_suppliers", "max_nomenclature_per_supplier", "max_manufacturers", "max_customers")
    )
    # Счётчики использования для отображения в UI (N/limit)
    usage = {}
    for model, key, deleted_attr in [
        (Nomenclature, "nomenclature", "is_deleted"),
        (Supplier, "suppliers", "is_deleted"),
        (Manufacturer, "manufacturers", "is_deleted"),
        (Customer, "customers", "is_deleted"),
    ]:
        q = select(func.count()).select_from(model).where(model.company_id == company_id)
        if deleted_attr and hasattr(model, deleted_attr):
            q = q.where(getattr(model, deleted_attr) == False)
        usage[key] = (await db.execute(q)).scalar() or 0
    return {
        "status": sub.status,
        "plan_name": plan.name,
        "expires_at": sub.expires_at.isoformat() if sub.expires_at else None,
        "is_demo": is_demo,
        "limits": {
            "max_nomenclature": getattr(plan, "max_nomenclature", None),
            "max_suppliers": getattr(plan, "max_suppliers", None),
            "max_nomenclature_per_supplier": getattr(plan, "max_nomenclature_per_supplier", None),
            "max_manufacturers": getattr(plan, "max_manufacturers", None),
            "max_customers": getattr(plan, "max_customers", None),
        } if has_limits else None,
        "usage": usage,
        "limits_note": (
            f"До {getattr(plan, 'max_suppliers', 3)} поставщиков, до {getattr(plan, 'max_nomenclature_per_supplier', 50)} товаров по каждому поставщику. Остальной функционал без ограничений."
            if getattr(plan, "plan_type", None) == "trader" else None
        ),
    }


@router.get("/invoices")
async def list_invoices(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """История счётов/оплат компании."""
    uc = await db.execute(select(UserCompany).where(UserCompany.user_id == user.id).limit(1))
    row = uc.scalar_one_or_none()
    if not row:
        return {"items": []}
    result = await db.execute(
        select(Invoice)
        .where(Invoice.company_id == row.company_id)
        .order_by(Invoice.created_at.desc())
        .limit(50)
    )
    invoices = result.scalars().all()
    return {
        "items": [
            {
                "id": inv.id,
                "invoice_number": inv.invoice_number,
                "amount": inv.amount,
                "status": inv.status,
                "period_start": inv.period_start.isoformat() if inv.period_start else None,
                "period_end": inv.period_end.isoformat() if inv.period_end else None,
                "created_at": inv.created_at.isoformat() if inv.created_at else None,
            }
            for inv in invoices
        ]
    }


class SaveQrCodeRequest(BaseModel):
    entity_type: str
    entity_id: int
    name: str
    qr_url: str
    image_base64: str


@router.post("/qr-codes")
async def save_qr_code(
    data: SaveQrCodeRequest,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Сохранить сгенерированный QR-код."""
    company_id = await get_user_company_id(user, db)
    if not company_id:
        raise HTTPException(403, "Нет компании")
    import base64
    raw = base64.b64decode(data.image_base64.split(",", 1)[-1] if "," in data.image_base64 else data.image_base64)
    if len(raw) > 2 * 1024 * 1024:
        raise HTTPException(400, "Размер изображения до 2 МБ")
    root = _storage_root()
    from time import time
    rel = f"qr_codes/company_{company_id}/{data.entity_type}_{data.entity_id}_{int(time() * 1000)}.png"
    path = root / rel
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "wb") as f:
        f.write(raw)
    db.add(SavedQrCode(
        company_id=company_id,
        entity_type=data.entity_type,
        entity_id=data.entity_id,
        name=data.name,
        qr_url=data.qr_url,
        file_path=rel,
    ))
    await db.commit()
    return {"ok": True}


@router.get("/qr-codes")
async def list_qr_codes(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Список сохранённых QR-кодов."""
    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        return {"items": []}
    result = await db.execute(
        select(SavedQrCode)
        .where(SavedQrCode.company_id.in_(company_ids))
        .order_by(SavedQrCode.created_at.desc())
        .limit(100)
    )
    items = result.scalars().all()
    return {
        "items": [
            {
                "id": q.id,
                "entity_type": q.entity_type,
                "entity_id": q.entity_id,
                "name": q.name,
                "qr_url": q.qr_url,
                "download_url": f"/api/cabinet/qr-codes/{q.id}/download",
                "created_at": q.created_at.isoformat(),
            }
            for q in items
        ]
    }


@router.get("/qr-codes/{qr_id}/download")
async def download_qr_code(
    qr_id: int,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Скачать сохранённый QR-код."""
    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        raise HTTPException(404, "Не найдено")
    result = await db.execute(
        select(SavedQrCode).where(
            SavedQrCode.id == qr_id,
            SavedQrCode.company_id.in_(company_ids),
        )
    )
    qr = result.scalar_one_or_none()
    if not qr:
        raise HTTPException(404, "QR-код не найден")
    root = _storage_root()
    path = root / qr.file_path
    if not path.exists():
        raise HTTPException(404, "Файл не найден")
    return FileResponse(path, media_type="image/png", filename=f"qr-{qr.entity_type}-{qr.entity_id}.png")


class DeleteQrCodesRequest(BaseModel):
    ids: list[int]


@router.delete("/qr-codes")
async def delete_qr_codes(
    data: DeleteQrCodesRequest,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Удалить выбранные QR-коды."""
    if not data.ids:
        return {"ok": True}
    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        raise HTTPException(403, "Нет доступа")
    result = await db.execute(
        select(SavedQrCode).where(
            SavedQrCode.id.in_(data.ids),
            SavedQrCode.company_id.in_(company_ids),
        )
    )
    items = result.scalars().all()
    for q in items:
        root = _storage_root()
        path = root / q.file_path
        if path.exists():
            try:
                path.unlink()
            except OSError:
                pass
        await db.delete(q)
    await db.commit()
    return {"ok": True}


# --- API Keys (Интеграции) ---


class CreateApiKeyRequest(BaseModel):
    name: str
    company_id: int | None = None
    scope: str = "read,write"
    expires_at: str | None = None  # ISO date или None


@router.get("/api-keys")
async def list_api_keys(
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Список API-ключей компании (для интеграций)."""
    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        return {"items": []}
    r = await db.execute(
        select(ApiKey)
        .where(ApiKey.company_id.in_(company_ids), ApiKey.is_active == True)
        .order_by(ApiKey.created_at.desc())
    )
    keys = r.scalars().all()
    return {
        "items": [
            {
                "id": k.id,
                "name": k.name,
                "scope": getattr(k, "scope", "read,write"),
                "expires_at": k.expires_at.isoformat() if getattr(k, "expires_at", None) else None,
                "last_used_at": k.last_used_at.isoformat() if k.last_used_at else None,
                "created_at": k.created_at.isoformat() if k.created_at else None,
            }
            for k in keys
        ]
    }


@router.post("/api-keys")
async def create_api_key(
    body: CreateApiKeyRequest,
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Создание API-ключа. Ключ показывается один раз — сохраните его."""
    cid = body.company_id or await get_user_company_id(user, db)
    if not cid:
        raise HTTPException(400, "Не указана компания (company_id)")
    scope = (body.scope or "read,write").strip() or "read,write"
    expires_at = None
    if body.expires_at:
        try:
            from datetime import datetime
            expires_at = datetime.fromisoformat(body.expires_at.replace("Z", "+00:00"))
        except (ValueError, TypeError):
            pass
    raw_key = f"ikam_{secrets.token_urlsafe(32)}"
    key_hash = hash_api_key(raw_key)
    api_key = ApiKey(company_id=cid, name=body.name or "API-ключ", key_hash=key_hash, scope=scope, expires_at=expires_at, created_by=user.id)
    db.add(api_key)
    await db.commit()
    await db.refresh(api_key)
    return {
        "id": api_key.id,
        "name": api_key.name,
        "key": raw_key,
        "message": "Сохраните ключ — он больше не будет показан",
    }


@router.delete("/api-keys/{key_id}")
async def revoke_api_key(
    key_id: int,
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Отзыв (деактивация) API-ключа."""
    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        raise HTTPException(403, "Нет доступа")
    r = await db.execute(select(ApiKey).where(ApiKey.id == key_id, ApiKey.company_id.in_(company_ids)))
    row = r.scalar_one_or_none()
    if not row:
        raise HTTPException(404, "Ключ не найден")
    row.is_active = False
    await db.commit()
    return {"ok": True}


# --- Webhooks ---


WEBHOOK_EVENTS = [
    "nomenclature.created", "nomenclature.updated", "nomenclature.deleted",
    "supply.created", "supply.updated",
    "supplier.created", "supplier.updated",
    "contract.created", "contract.updated",
]


class CreateWebhookRequest(BaseModel):
    name: str
    url: str
    events: list[str] | None = None


@router.get("/webhooks")
async def list_webhooks(
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Список webhooks для push-уведомлений."""
    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        return {"items": []}
    r = await db.execute(
        select(Webhook).where(Webhook.company_id.in_(company_ids)).order_by(Webhook.created_at.desc())
    )
    hooks = r.scalars().all()
    return {
        "items": [
            {
                "id": h.id,
                "name": h.name,
                "url": h.url,
                "events": h.events or [],
                "is_active": h.is_active,
                "last_triggered_at": h.last_triggered_at.isoformat() if h.last_triggered_at else None,
            }
            for h in hooks
        ]
    }


@router.post("/webhooks")
async def create_webhook(
    body: CreateWebhookRequest,
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Создание webhook."""
    cid = await get_user_company_id(user, db)
    if not cid:
        raise HTTPException(400, "Укажите компанию")
    if not body.url.startswith("https://") and not body.url.startswith("http://"):
        raise HTTPException(400, "URL должен начинаться с https:// или http://")
    events = [e for e in (body.events or []) if e in WEBHOOK_EVENTS] or WEBHOOK_EVENTS[:3]
    wh = Webhook(company_id=cid, name=body.name or "Webhook", url=body.url, events=events)
    db.add(wh)
    await db.commit()
    await db.refresh(wh)
    return {"id": wh.id, "name": wh.name, "url": wh.url, "events": wh.events}


@router.delete("/webhooks/{webhook_id}")
async def delete_webhook(
    webhook_id: int,
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Удаление webhook."""
    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        raise HTTPException(403, "Нет доступа")
    r = await db.execute(select(Webhook).where(Webhook.id == webhook_id, Webhook.company_id.in_(company_ids)))
    row = r.scalar_one_or_none()
    if not row:
        raise HTTPException(404, "Webhook не найден")
    await db.delete(row)
    await db.commit()
    return {"ok": True}


# --- Support Tickets ---

MAX_ATTACHMENT_BYTES = 10 * 1024 * 1024  # 10 МБ
ALLOWED_ATTACHMENT_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".txt", ".zip"}


def _safe_filename(name: str) -> str:
    """Безопасное имя файла."""
    base = Path(name).stem[:80] if name else "file"
    ext = Path(name).suffix.lower()[:20] if name else ""
    if ext not in ALLOWED_ATTACHMENT_EXTS and ext:
        ext = ".dat"
    elif not ext:
        ext = ".bin"
    return f"{base}{ext}".replace(" ", "_")


async def _save_attachment(file: UploadFile, subdir: str) -> tuple[str, str, int]:
    """Сохраняет файл, возвращает (storage_path, filename, size)."""
    content = await file.read()
    if len(content) > MAX_ATTACHMENT_BYTES:
        raise HTTPException(400, f"Файл до 10 МБ. Размер: {len(content) // (1024*1024)} МБ")
    ext = Path(file.filename or "").suffix.lower()
    if ext and ext not in ALLOWED_ATTACHMENT_EXTS:
        raise HTTPException(400, f"Разрешены: {', '.join(ALLOWED_ATTACHMENT_EXTS)}")
    root = _storage_root()
    rel_dir = f"support/{subdir}"
    (root / rel_dir).mkdir(parents=True, exist_ok=True)
    safe = secrets.token_hex(8) + "_" + _safe_filename(file.filename or "file")
    rel = f"{rel_dir}/{safe}"
    path = root / rel
    with open(path, "wb") as f:
        f.write(content)
    return rel, file.filename or safe, len(content)


@router.get("/tickets")
async def list_tickets(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Список тикетов пользователя (админ видит все тикеты своей компании)."""
    company_ids = await get_user_company_ids(user, db)
    q = select(SupportTicket).order_by(SupportTicket.updated_at.desc())
    if company_ids is not None:
        if not company_ids:
            return {"items": []}
        q = q.where(SupportTicket.company_id.in_(company_ids))
    r = await db.execute(q)
    tickets = r.scalars().all()
    return {
        "items": [
            {
                "id": t.id,
                "subject": t.subject,
                "status": t.status,
                "created_at": t.created_at.isoformat(),
                "updated_at": t.updated_at.isoformat(),
            }
            for t in tickets
        ]
    }


@router.post("/tickets")
async def create_ticket(
    subject: str = Form(...),
    body: str = Form(...),
    files: list[UploadFile] = File(default=[]),
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Создать тикет. Файлы до 10 МБ."""
    cid = await get_user_company_id(user, db)
    if not cid:
        raise HTTPException(400, "Нет компании")
    t = SupportTicket(user_id=user.id, company_id=cid, subject=subject.strip(), body=body.strip())
    db.add(t)
    await db.flush()
    for f in files or []:
        if f.filename and not f.filename.startswith("."):
            try:
                rel, fname, sz = await _save_attachment(f, f"tickets/{t.id}")
                db.add(TicketAttachment(ticket_id=t.id, reply_id=None, storage_path=rel, filename=fname, size_bytes=sz))
            except HTTPException:
                raise
            except Exception as e:
                logging.warning("ticket attachment save failed: %s", e)
    await db.commit()
    await db.refresh(t)
    return {"id": t.id, "subject": t.subject, "status": t.status}


@router.get("/tickets/attachment/{att_id}")
async def get_ticket_attachment(
    att_id: int,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Скачать вложение тикета."""
    r = await db.execute(
        select(TicketAttachment, SupportTicket).join(SupportTicket, TicketAttachment.ticket_id == SupportTicket.id).where(TicketAttachment.id == att_id)
    )
    row = r.first()
    if not row:
        raise HTTPException(404, "Вложение не найдено")
    att, tick = row[0], row[1]
    company_ids = await get_user_company_ids(user, db)
    is_staff = user.role in (Role.admin, Role.super_admin)
    if not is_staff and (not company_ids or tick.company_id not in company_ids):
        raise HTTPException(403, "Нет доступа")
    root = _storage_root()
    path = root / att.storage_path
    if not path.exists():
        raise HTTPException(404, "Файл не найден")
    return FileResponse(path, filename=att.filename, media_type="application/octet-stream")


@router.get("/tickets/{tid}")
async def get_ticket(
    tid: int,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Получить тикет с ответами."""
    company_ids = await get_user_company_ids(user, db)
    r = await db.execute(select(SupportTicket).where(SupportTicket.id == tid))
    t = r.scalar_one_or_none()
    if not t:
        raise HTTPException(404, "Тикет не найден")
    is_staff = user.role in (Role.admin, Role.super_admin)
    if not is_staff and (company_ids is None or t.company_id not in company_ids):
        raise HTTPException(403, "Нет доступа")
    replies_r = await db.execute(
        select(TicketReply).where(TicketReply.ticket_id == tid).order_by(TicketReply.created_at.asc())
    )
    replies = replies_r.scalars().all()
    att_r = await db.execute(select(TicketAttachment).where(TicketAttachment.ticket_id == tid))
    att_list = att_r.scalars().all()
    att_by_ticket = [{"id": a.id, "filename": a.filename, "size_bytes": a.size_bytes} for a in att_list if a.reply_id is None]
    att_by_reply: dict[int, list] = {}
    for a in att_list:
        if a.reply_id is not None:
            att_by_reply.setdefault(a.reply_id, []).append({"id": a.id, "filename": a.filename, "size_bytes": a.size_bytes})
    return {
        "id": t.id,
        "subject": t.subject,
        "body": t.body,
        "status": t.status,
        "created_at": t.created_at.isoformat(),
        "updated_at": t.updated_at.isoformat(),
        "attachments": att_by_ticket,
        "replies": [
            {"id": r.id, "body": r.body, "is_staff": r.is_staff, "created_at": r.created_at.isoformat(), "attachments": att_by_reply.get(r.id, [])}
            for r in replies
        ],
    }


@router.post("/tickets/{tid}/reply")
async def add_ticket_reply(
    tid: int,
    body: str = Form(...),
    files: list[UploadFile] = File(default=[]),
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Добавить ответ к тикету. Файлы до 10 МБ."""
    r = await db.execute(select(SupportTicket).where(SupportTicket.id == tid))
    t = r.scalar_one_or_none()
    if not t:
        raise HTTPException(404, "Тикет не найден")
    company_ids = await get_user_company_ids(user, db)
    is_staff = user.role in (Role.admin, Role.super_admin)
    if not is_staff:
        if not company_ids or t.company_id not in company_ids:
            raise HTTPException(403, "Нет доступа")
        if t.user_id != user.id:
            raise HTTPException(403, "Можно отвечать только в своих тикетах")
    reply = TicketReply(ticket_id=tid, author_id=user.id, body=body.strip(), is_staff=is_staff)
    db.add(reply)
    await db.flush()
    for f in files or []:
        if f.filename and not f.filename.startswith("."):
            try:
                rel, fname, sz = await _save_attachment(f, f"tickets/{tid}/replies")
                db.add(TicketAttachment(ticket_id=tid, reply_id=reply.id, storage_path=rel, filename=fname, size_bytes=sz))
            except HTTPException:
                raise
            except Exception as e:
                logging.warning("ticket reply attachment save failed: %s", e)
    if is_staff:
        n = Notification(user_id=t.user_id, title="Ответ по тикету", body=f"Тикет «{t.subject}»: новый ответ от поддержки", type="info", link=f"/cabinet/tickets?id={tid}")
        db.add(n)
    await db.commit()
    await db.refresh(reply)
    return {"id": reply.id, "body": reply.body, "is_staff": reply.is_staff, "created_at": reply.created_at.isoformat()}


# --- Suggestions ---


@router.get("/suggestions")
async def list_suggestions(
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Список предложений пользователя."""
    company_ids = await get_user_company_ids(user, db)
    q = select(Suggestion).order_by(Suggestion.updated_at.desc())
    if company_ids is not None:
        if not company_ids:
            return {"items": []}
        q = q.where(Suggestion.company_id.in_(company_ids))
    r = await db.execute(q)
    items = r.scalars().all()
    return {
        "items": [
            {"id": s.id, "subject": s.subject, "status": s.status, "created_at": s.created_at.isoformat()}
            for s in items
        ]
    }


@router.post("/suggestions")
async def create_suggestion(
    subject: str = Form(...),
    body: str = Form(...),
    files: list[UploadFile] = File(default=[]),
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Создать предложение по улучшению. Файлы до 10 МБ."""
    cid = await get_user_company_id(user, db)
    if not cid:
        raise HTTPException(400, "Нет компании")
    s = Suggestion(user_id=user.id, company_id=cid, subject=subject.strip(), body=body.strip(), status="received")
    db.add(s)
    await db.flush()
    for f in files or []:
        if f.filename and not f.filename.startswith("."):
            try:
                rel, fname, sz = await _save_attachment(f, f"suggestions/{s.id}")
                db.add(SuggestionAttachment(suggestion_id=s.id, update_id=None, storage_path=rel, filename=fname, size_bytes=sz))
            except HTTPException:
                raise
            except Exception as e:
                logging.warning("suggestion attachment save failed: %s", e)
    up = SuggestionUpdate(suggestion_id=s.id, body="Ваше предложение получено. Спасибо!", status_after="received")
    db.add(up)
    n = Notification(user_id=user.id, title="Предложение получено", body=f"«{s.subject}» — получено и принято к учёту.", type="success", link=f"/cabinet/suggestions?id={s.id}")
    db.add(n)
    await db.commit()
    await db.refresh(s)
    return {"id": s.id, "subject": s.subject, "status": s.status}


@router.get("/suggestions/attachment/{att_id}")
async def get_suggestion_attachment(
    att_id: int,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Скачать вложение предложения."""
    r = await db.execute(
        select(SuggestionAttachment, Suggestion).join(Suggestion, SuggestionAttachment.suggestion_id == Suggestion.id).where(SuggestionAttachment.id == att_id)
    )
    row = r.first()
    if not row:
        raise HTTPException(404, "Вложение не найдено")
    att, sug = row[0], row[1]
    company_ids = await get_user_company_ids(user, db)
    is_staff = user.role in (Role.admin, Role.super_admin)
    if not is_staff and (not company_ids or sug.company_id not in company_ids):
        raise HTTPException(403, "Нет доступа")
    root = _storage_root()
    path = root / att.storage_path
    if not path.exists():
        raise HTTPException(404, "Файл не найден")
    return FileResponse(path, filename=att.filename, media_type="application/octet-stream")


@router.get("/suggestions/{sid}")
async def get_suggestion(
    sid: int,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Получить предложение с обновлениями."""
    company_ids = await get_user_company_ids(user, db)
    r = await db.execute(select(Suggestion).where(Suggestion.id == sid))
    s = r.scalar_one_or_none()
    if not s:
        raise HTTPException(404, "Не найдено")
    is_staff = user.role in (Role.admin, Role.super_admin)
    if not is_staff and (company_ids is None or s.company_id not in company_ids):
        raise HTTPException(403, "Нет доступа")
    upd_r = await db.execute(
        select(SuggestionUpdate).where(SuggestionUpdate.suggestion_id == sid).order_by(SuggestionUpdate.created_at.asc())
    )
    updates = upd_r.scalars().all()
    att_r = await db.execute(select(SuggestionAttachment).where(SuggestionAttachment.suggestion_id == sid))
    att_list = att_r.scalars().all()
    att_main = [{"id": a.id, "filename": a.filename, "size_bytes": a.size_bytes} for a in att_list if a.update_id is None]
    att_by_update: dict[int, list] = {}
    for a in att_list:
        if a.update_id is not None:
            att_by_update.setdefault(a.update_id, []).append({"id": a.id, "filename": a.filename, "size_bytes": a.size_bytes})
    return {
        "id": s.id,
        "subject": s.subject,
        "body": s.body,
        "status": s.status,
        "created_at": s.created_at.isoformat(),
        "attachments": att_main,
        "updates": [
            {"id": u.id, "body": u.body, "status_after": u.status_after, "created_at": u.created_at.isoformat(), "attachments": att_by_update.get(u.id, [])}
            for u in updates
        ],
    }


@router.post("/suggestions/{sid}/update")
async def add_suggestion_update(
    sid: int,
    body: str = Form(...),
    status: str | None = Form(None),
    files: list[UploadFile] = File(default=[]),
    user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Добавить обновление к предложению (только для админов). Файлы до 10 МБ."""
    r = await db.execute(select(Suggestion).where(Suggestion.id == sid))
    s = r.scalar_one_or_none()
    if not s:
        raise HTTPException(404, "Не найдено")
    prev = s.status
    if status:
        s.status = status
    up = SuggestionUpdate(suggestion_id=sid, author_id=user.id, body=body.strip(), status_before=prev, status_after=status or prev)
    db.add(up)
    await db.flush()
    for f in files or []:
        if f.filename and not f.filename.startswith("."):
            try:
                rel, fname, sz = await _save_attachment(f, f"suggestions/{sid}/updates")
                db.add(SuggestionAttachment(suggestion_id=sid, update_id=up.id, storage_path=rel, filename=fname, size_bytes=sz))
            except HTTPException:
                raise
            except Exception as e:
                logging.warning("suggestion update attachment save failed: %s", e)
    n = Notification(user_id=s.user_id, title="Обновление по предложению", body=body, type="info", link=f"/cabinet/suggestions?id={sid}")
    db.add(n)
    await db.commit()
    return {"id": up.id, "body": up.body, "status_after": up.status_after, "created_at": up.created_at.isoformat()}
