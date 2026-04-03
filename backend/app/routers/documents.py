# -*- coding: utf-8 -*-
"""Загрузка и скачивание документов."""

import io
import json
import logging
import os
import re
import secrets
import zipfile
from datetime import date, datetime
from pathlib import Path
from typing import Annotated

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query, Request
from fastapi.responses import FileResponse, Response
from pypdf import PdfReader
from sqlalchemy import select, func, or_, and_
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.database import get_db
from app.dependencies import require_trial_or_subscription, get_user_company_ids, get_user_company_id, require_can_delete_documents, check_entity_limit, check_nomenclature_per_supplier_limit
from app.models.user import User
from app.models.document import Document, ENTITY_TYPES
from app.models.entity import Nomenclature, Supplier, Manufacturer, Supply, Contract
from app.models.import_job import ImportJob

router = APIRouter()
logger = logging.getLogger(__name__)


def _cabinet_payment_subscription_url() -> str | None:
    """Ссылка на оплату в кабинете — только если включена ЮKassa."""
    return "/cabinet/payment" if get_settings().feature_yookassa else None

# Разрешённые MIME для загрузки
ALLOWED_MIMES = {
    "application/pdf",
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/gif",
    "image/webp",
    "application/msword",  # .doc
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/vnd.ms-excel",  # .xls
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # .xlsx
    "application/zip",
}


def _storage_root() -> Path:
    """Корень хранилища документов."""
    p = get_settings().storage_path
    if not p.is_absolute():
        p = Path(__file__).resolve().parents[1] / "storage"
    p.mkdir(parents=True, exist_ok=True)
    return p


def _safe_filename(name: str) -> str:
    """Безопасное имя файла."""
    base = os.path.basename(name)
    if not base:
        base = "document"
    # убираем путь
    safe = "".join(c for c in base if c.isalnum() or c in "._- ")
    return safe[:200] or "document"


ALLOWED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".gif", ".webp", ".doc", ".docx", ".xls", ".xlsx", ".zip"}


def _is_allowed_file(mime: str, ext: str) -> bool:
    ext_lower = (ext or "").lower()
    return mime in ALLOWED_MIMES or ext_lower in ALLOWED_EXTENSIONS or mime.startswith("image/")


# Magic bytes для проверки реального типа файла (защита от подделки MIME)
_MAGIC_SIGNATURES: dict[str, list[bytes]] = {
    "pdf": [b"%PDF"],
    "jpeg": [b"\xff\xd8\xff"],
    "png": [b"\x89PNG\r\n\x1a\n"],
    "gif": [b"GIF87a", b"GIF89a"],
    "webp": [b"RIFF"],  # RIFF....WEBP, проверяем RIFF
    "zip": [b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08"],
    "docx": [b"PK\x03\x04"],  # Office Open XML — ZIP
    "xlsx": [b"PK\x03\x04"],
    "doc": [b"\xd0\xcf\x11\xe0"],  # OLE (MS Office binary)
    "xls": [b"\xd0\xcf\x11\xe0"],
}

_EXT_TO_MAGIC_KEYS: dict[str, list[str]] = {
    ".pdf": ["pdf"],
    ".jpg": ["jpeg"], ".jpeg": ["jpeg"],
    ".png": ["png"],
    ".gif": ["gif"],
    ".webp": ["webp"],
    ".zip": ["zip"],
    ".docx": ["docx", "zip"],
    ".xlsx": ["xlsx", "zip"],
    ".doc": ["doc"],
    ".xls": ["xls"],
}


def _verify_magic_bytes(content: bytes, mime: str, ext: str) -> None:
    """Проверка magic bytes. Вызывает HTTPException при несоответствии."""
    if len(content) < 12:
        return
    ext_lower = (ext or "").lower()
    keys = _EXT_TO_MAGIC_KEYS.get(ext_lower, [])
    if not keys:
        if mime.startswith("image/"):
            keys = ["jpeg", "png", "gif", "webp"]
        elif "pdf" in (mime or ""):
            keys = ["pdf"]
        elif "sheet" in (mime or "") or "excel" in (mime or ""):
            keys = ["xlsx", "zip"]
        elif "word" in (mime or ""):
            keys = ["docx", "zip"]
        elif "zip" in (mime or ""):
            keys = ["zip"]
    if not keys:
        return
    for k in keys:
        sigs = _MAGIC_SIGNATURES.get(k, [])
        for sig in sigs:
            if content[:len(sig)] == sig:
                if k == "webp" and b"WEBP" not in content[:20]:
                    continue
                return
    raise HTTPException(400, f"Файл не соответствует заявленному типу ({ext or mime})")


def _check_file_size(content: bytes) -> None:
    """Проверка размера файла. Вызывает HTTPException при превышении лимита."""
    settings = get_settings()
    limit_bytes = settings.max_upload_size_mb * 1024 * 1024
    if len(content) > limit_bytes:
        raise HTTPException(
            413,
            f"Файл слишком большой. Максимум {settings.max_upload_size_mb} МБ.",
        )


async def _save_document(
    content: bytes,
    filename: str,
    mime: str,
    entity_type: str,
    entity_id: int,
    company_id: int,
    created_by: int,
    db: AsyncSession,
) -> dict:
    """Сохранить документ на диск и в БД. Возвращает данные созданного документа."""
    root = _storage_root()
    rel = f"{company_id}/{entity_type}/{entity_id}"
    dir_path = root / rel
    dir_path.mkdir(parents=True, exist_ok=True)
    ext = Path(filename or "").suffix or ".bin"
    unique = secrets.token_hex(8)
    safe = _safe_filename(filename or "doc") + f"_{unique}{ext}"
    storage_path = str(dir_path / safe)
    with open(storage_path, "wb") as f:
        f.write(content)
    rel_storage = f"{rel}/{safe}"
    doc = Document(
        company_id=company_id,
        entity_type=entity_type,
        entity_id=entity_id,
        filename=filename or safe,
        mime_type=mime,
        storage_path=rel_storage,
        file_size=len(content),
        created_by=created_by,
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)
    return {
        "id": doc.id,
        "filename": doc.filename,
        "mime_type": doc.mime_type,
        "file_size": doc.file_size,
        "entity_type": entity_type,
        "entity_id": entity_id,
    }


def _extract_files_from_zip(zip_content: bytes, zip_filename: str) -> list[tuple[str, bytes, str]]:
    """Извлечь файлы из ZIP. Возвращает [(filename, content, mime), ...]. Пропускает недопустимые."""
    result = []
    try:
        with zipfile.ZipFile(io.BytesIO(zip_content), "r") as zf:
            for name in zf.namelist():
                if name.endswith("/") or "__MACOSX" in name or name.startswith("."):
                    continue
                base = os.path.basename(name)
                if not base:
                    continue
                ext = Path(base).suffix.lower()
                mime_map = {
                    ".pdf": "application/pdf",
                    ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg",
                    ".png": "image/png",
                    ".gif": "image/gif",
                    ".webp": "image/webp",
                    ".doc": "application/msword",
                    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ".xls": "application/vnd.ms-excel",
                    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                }
                mime = mime_map.get(ext, "application/octet-stream")
                if ext not in (e for e in ALLOWED_EXTENSIONS if e != ".zip"):
                    continue
                try:
                    content = zf.read(name)
                except (zipfile.BadZipFile, KeyError, RuntimeError):
                    continue
                if len(content) > get_settings().max_upload_size_mb * 1024 * 1024:
                    raise HTTPException(413, f"Файл в архиве слишком большой: {base}")
                result.append((base, content, mime))
    except zipfile.BadZipFile:
        raise HTTPException(400, f"Неверный ZIP-архив: {zip_filename}")
    return result


def _extract_text_from_file(content: bytes, filename: str, mime: str) -> str:
    """Извлечь текст из PDF, DOCX, XLSX. Остальные типы возвращают пустую строку."""
    ext = Path(filename or "").suffix.lower()
    text_parts = []
    try:
        if mime == "application/pdf" or ext == ".pdf":
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
        elif mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" or ext == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    vals = [str(v) for v in (row or []) if v is not None and str(v).strip()]
                    if vals:
                        text_parts.append(" ".join(vals))
        elif mime == "application/msword" or ext == ".doc":
            pass
        elif mime == "application/vnd.ms-excel" or ext == ".xls":
            pass
    except Exception:
        pass
    return "\n".join(text_parts)


def _get_max_import_rows() -> int:
    return getattr(get_settings(), "max_import_rows", 100_000) or 100_000


def _extract_universal(content: bytes, filename: str, mime: str) -> dict:
    """Универсальное извлечение: все колонки и строки. Возвращает detected_columns для сопоставления полей."""
    max_rows = _get_max_import_rows()
    ext = Path(filename or "").suffix.lower()
    detected_columns: list[dict] = []
    header_dict: dict[str, str] = {}
    nomenclature_items: list[dict] = []
    rows_data: list[dict[int, str]] = []
    raw_text = ""

    def _cell_val(v) -> str:
        if v is None:
            return ""
        s = str(v).strip().replace("\n", " ")
        return s[:500] if len(s) > 500 else s

    try:
        if ext == ".xlsx" or (mime or "").endswith("sheet"):
            from openpyxl import load_workbook

            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            all_columns: dict[int, list[str]] = {}
            header_row_idx = -1
            max_cols = 0

            best_sheet = None
            best_header_row = -1
            best_score = 0

            for sheet in wb.worksheets:
                sheet_max = getattr(sheet, "max_row", None) or 0
                read_up_to = min(sheet_max or 1, max_rows + 100)
                rows_list = [list(r or []) for r in sheet.iter_rows(min_row=1, max_row=read_up_to, values_only=True)]
                if len(rows_list) < 2:
                    continue
                header_patterns = ("номер", "места", "наименование", "кол-во", "шт", "см", "маркировка", "страна", "тип", "грузового", "отгружаемой", "марка", "детали", "упаковк", "длина", "ширина", "высота", "вес", "производитель")
                data_like = ("ооо ", '"', "инн", "тел.", "адрес", "россия", "москва", "согласно контракту", "покупатель", "грузополучатель", "велесстрой")
                for ri, row in enumerate(rows_list[:50]):
                    row = list(row or [])
                    non_empty = [(i, _cell_val(v)) for i, v in enumerate(row) if v is not None and _cell_val(v)]
                    if len(non_empty) < 3:
                        continue
                    first_val = non_empty[0][1] if non_empty else ""
                    looks_like_data = first_val.replace(".", "").replace(",", "").replace(" ", "").replace("\xa0", "").isdigit() and len(first_val) <= 8
                    if looks_like_data:
                        continue
                    if any(d in first_val.lower()[:50] for d in data_like):
                        continue
                    header_cells = sum(2 if any(p in val.lower() for p in header_patterns) else (1 if len(val) > 5 else 0) for _, val in non_empty)
                    if header_cells >= 5 and header_cells > best_score:
                        best_score = header_cells
                        best_header_row = ri
                        best_sheet = rows_list

            if best_sheet is not None and best_header_row >= 0:
                rows_list = best_sheet
                header_row_idx = best_header_row
                header_row = list(rows_list[header_row_idx] or [])
                for ci, v in enumerate(header_row):
                    h = _cell_val(v)
                    if h:
                        all_columns[ci] = [h]

            if not all_columns:
                for sheet in wb.worksheets:
                    rows_list = [list(r or []) for r in sheet.iter_rows(values_only=True)]
                    if rows_list:
                        for ci, v in enumerate(rows_list[0]):
                            h = _cell_val(v)
                            if h:
                                all_columns[ci] = [h]
                        header_row_idx = 0
                        best_sheet = rows_list
                        break

            data_rows = best_sheet if best_sheet else []
            if not data_rows and all_columns:
                for sheet in wb.worksheets:
                    data_rows = [list(r or []) for r in sheet.iter_rows(values_only=True)]
                    if len(data_rows) >= 2:
                        break
            start = header_row_idx + 1 if header_row_idx >= 0 else 1
            end = min(len(data_rows), start + max_rows)
            rows_data: list[dict[int, str]] = []
            for ri in range(start, end):
                row = data_rows[ri]
                row_str = " ".join(str(v or "").lower() for v in (row or []))
                if "итого" in row_str or "total" in row_str or "подпись" in row_str:
                    break
                row_dict: dict[int, str] = {}
                has_any = False
                for ci in list(all_columns.keys()):
                    val = _cell_val(row[ci]) if ci < len(row) else ""
                    skip_vals = ("не применимо", "not applicable", "n/a", "-", "")
                    if val and str(val).strip().lower() not in skip_vals:
                        has_any = True
                        row_dict[ci] = str(val).strip()
                        if len(all_columns[ci]) < 6:
                            if val not in all_columns[ci]:
                                all_columns[ci].append(val)
                if has_any:
                    rows_data.append(row_dict)

            for ci in sorted(all_columns.keys()):
                vals = all_columns[ci]
                if vals:
                    name = vals[0]
                    example = next((v for v in vals[1:] if v and len(v) > 1), vals[0])
                    detected_columns.append({
                        "index": ci,
                        "name": name[:200].replace("\n", " ").strip(),
                        "example": example[:200],
                        "samples": vals[1:6],
                    })
            if best_sheet is None:
                rows_data = []
            wb.close()

        elif ext == ".docx" or (mime or "").endswith("wordprocessingml"):
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(content))
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    raw_text += t + "\n"
            for table in doc.tables:
                rows = [[_cell_val(c.text) for c in row.cells] for row in table.rows]
                if len(rows) >= 2:
                    headers = rows[0]
                    for ci, h in enumerate(headers):
                        if h:
                            examples = [r[ci] if ci < len(r) and r[ci] else "" for r in rows[1:6]]
                            examples = [e for e in examples if e]
                            detected_columns.append({
                                "index": ci,
                                "name": h[:200],
                                "example": examples[0] if examples else h[:200],
                                "samples": examples[:5],
                            })
                    for r in rows[1:]:
                        row_dict: dict[int, str] = {ci: (r[ci] if ci < len(r) and r[ci] else "") for ci in range(len(headers)) if headers[ci]}
                        if any(row_dict.values()):
                            rows_data.append(row_dict)
            if not detected_columns and raw_text:
                for i, line in enumerate(raw_text.split("\n")[:20]):
                    if line.strip():
                        detected_columns.append({
                            "index": i,
                            "name": f"Строка {i + 1}",
                            "example": line[:200],
                            "samples": [],
                        })

        elif ext == ".pdf" or (mime or "").endswith("pdf"):
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    raw_text += t + "\n"
            lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
            for i, line in enumerate(lines[:30]):
                detected_columns.append({
                    "index": i,
                    "name": f"Блок {i + 1}",
                    "example": line[:200],
                    "samples": [],
                })
    except Exception:
        pass

    return {
        "detected_columns": detected_columns,
        "rows_data": rows_data,
        "raw_preview": raw_text[:1500] if raw_text else "",
        "header": header_dict,
        "nomenclature_items": nomenclature_items,
    }


def _build_nom_items_from_universal(rows_data: list[dict[int, str]], detected_columns: list[dict]) -> list[dict]:
    """Преобразовать rows_data в nomenclature_items по авто-маппингу колонок."""
    FIELD_MAP = [
        (["таговый", "tag number", "марка", "тип", "маркировка"], "tag_number"),
        (["артикул", "vendor code", "код изделия", "код товара"], "article"),
        (["наименование", "product", "изделие", "описание", "товар"], "name"),
        (["кол-во", "qnty", "quan", "количество", "шт", "единиц"], "quantity"),
        (["производитель", "manufacturer"], "manufacturer"),
        (["грузового места", "package", "номер места", "упаковк"], "package_number"),
    ]
    col_to_field: dict[int, str] = {}
    for dc in detected_columns:
        ci = dc.get("index", -1)
        name = (dc.get("name") or "").lower()
        if ci < 0:
            continue
        for keywords, field in FIELD_MAP:
            if any(kw in name for kw in keywords):
                if field not in col_to_field.values():
                    col_to_field[ci] = field
                break
    result: list[dict] = []
    for row_dict in rows_data:
        item: dict = {}
        for ci, field in col_to_field.items():
            val = row_dict.get(ci, "")
            if val and str(val).strip():
                item[field] = str(val).strip()
        if item.get("name") or item.get("tag_number") or item.get("article"):
            if "quantity" not in item:
                item["quantity"] = 1
            if item.get("article") and "code" not in item:
                item["code"] = item["article"]
            result.append(item)
    return result


def _parse_packing_document_xlsx(content: bytes, filename: str) -> dict | None:
    """Структурированное извлечение из XLSX (ведомость грузового места / packing note)."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        header: dict[str, str] = {}
        nomenclature_items: list[dict] = []

        def _cell(row: list, col: int) -> str:
            if col < len(row) and row[col] is not None:
                return str(row[col]).strip()
            return ""

        def _find_val(row: list, start: int = 3) -> str:
            for c in range(start, min(start + 6, len(row))):
                v = _cell(row, c)
                if v and v.lower() not in ("не применимо", "not applicable", "n/a"):
                    return v
            return ""

        def _find_in_row(rows: list, *substrs: str) -> tuple[int, str] | None:
            for ri, row in enumerate(rows):
                row_str = " ".join(str(v or "").lower() for v in (row or []))
                for s in substrs:
                    if s.lower() in row_str:
                        return ri, _find_val(list(row or []))
            return None

        for sheet in wb.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if len(rows) < 3:
                continue
            rows_list = [list(r or []) for r in rows]

            # Шапка: ищем метки и значения
            for ri, row in enumerate(rows_list):
                r0 = _cell(row, 1)
                r5 = _cell(row, 5)
                r6 = _cell(row, 6) if len(row) > 6 else ""
                if "ведомость" in r0.lower() or "packing note" in r0.lower():
                    header["номер_документа"] = r5 or r6 or header.get("номер_документа", "")
                if "договор" in r0.lower() and "contract" in r0.lower():
                    header["договор_и_дата"] = r5 or r6 or header.get("договор_и_дата", "")
                if "приложение" in r0.lower() and "appendix" in r0.lower():
                    header["приложение_к_договору"] = r5 or r6 or header.get("приложение_к_договору", "")
                if "заказная спецификация" in r0.lower():
                    header["заказная_спецификация"] = r5 or r6 or header.get("заказная_спецификация", "")
                if "базис поставки" in r0.lower() or "delivery terms" in r0.lower():
                    header["базис_поставки"] = r5 or r6 or header.get("базис_поставки", "")
                if "страна происхождения" in r0.lower():
                    header["страна_происхождения"] = _find_val(row, 8) or header.get("страна_происхождения", "")
                if "адрес отгрузки" in r0.lower() or "address of pickup" in r0.lower():
                    header["адрес_отгрузки"] = _find_val(row, 8) or header.get("адрес_отгрузки", "")
                if "наименование проекта" in r0.lower():
                    header["наименование_проекта"] = _find_val(row, 8) or header.get("наименование_проекта", "")
                if "адрес поставки" in r0.lower() or "delivery address" in r0.lower():
                    header["адрес_поставки"] = _find_val(row, 8) or header.get("адрес_поставки", "")
                if "наименование поставляемого" in r0.lower():
                    header["наименование_поставляемого_изделия"] = _find_val(row, 8) or header.get("наименование_поставляемого_изделия", "")
                if "продавец" in r0.lower() and "seller" in r0.lower():
                    nrow = rows_list[ri + 1] if ri + 1 < len(rows_list) else []
                    header["продавец"] = _find_val(nrow, 1) or header.get("продавец", "")
                if "покупатель" in r0.lower() and "buyer" in r0.lower():
                    nrow = rows_list[ri + 1] if ri + 1 < len(rows_list) else []
                    header["покупатель"] = _find_val(nrow, 8) or header.get("покупатель", "")
                if "отправитель" in r0.lower() and "shipper" in r0.lower():
                    nrow = rows_list[ri + 1] if ri + 1 < len(rows_list) else []
                    header["отправитель"] = _find_val(nrow, 1) or header.get("отправитель", "")
                if "грузополучатель" in r0.lower() and "consignee" in r0.lower():
                    nrow = rows_list[ri + 1] if ri + 1 < len(rows_list) else []
                    header["грузополучатель"] = _find_val(nrow, 8) or header.get("грузополучатель", "")

            # Таблица номенклатуры: ищем строку с заголовками
            tag_col = art_col = name_col = qty_col = mfr_col = country_col = package_col = -1
            data_start = 0
            for ri, row in enumerate(rows_list):
                row_str = " ".join(str(v or "").lower() for v in (row or []))
                if any(x in row_str for x in ("таговый", "tag number", "артикул", "наименование", "кол-во", "количество", "номер", "марка")):
                    for ci, v in enumerate(row or []):
                        vstr = str(v or "").lower()
                        if "таговый" in vstr or "tag number" in vstr or ("маркировка" in vstr and "номер" in vstr) or ("tag" in vstr and "no" in vstr):
                            tag_col = ci
                        if "артикул" in vstr or "vendor code" in vstr or "код" in vstr:
                            if art_col < 0:
                                art_col = ci
                        if ("наименование" in vstr and "товара" in vstr) or "product" in vstr or ("наименование" in vstr and "изделия" in vstr) or "описание" in vstr:
                            name_col = ci
                        if "кол-во" in vstr or "qnty" in vstr or "quan" in vstr or "количество" in vstr or "шт" in vstr or "единиц" in vstr:
                            qty_col = ci
                        if "производитель" in vstr or "manufacturer" in vstr:
                            mfr_col = ci
                        if ("страна" in vstr and "происхождения" in vstr) or "country" in vstr:
                            country_col = ci
                        if "грузового места" in vstr or ("unique identifier" in vstr and "package" in vstr) or "номер места" in vstr or "упаковк" in vstr:
                            package_col = ci
                    if tag_col >= 0 or art_col >= 0 or name_col >= 0:
                        data_start = ri + 2
                        break

            if tag_col < 0 and art_col < 0 and name_col < 0:
                for ri, row in enumerate(rows_list):
                    row_str = " ".join(str(v or "").lower() for v in (row or []))
                    if "item no" in row_str:
                        tag_col, art_col, name_col = 4, 5, 6
                        if len(row or []) > 10:
                            qty_col = 9
                        data_start = ri + 2
                        break
                if tag_col < 0 and art_col < 0 and name_col < 0:
                    for ri, row in enumerate(rows_list):
                        row_str = " ".join(str(v or "").lower() for v in (row or []))
                        if any(x in row_str for x in ("№ п/п", "позиция", "поз.", "наименование", "изделие", "ед.", "марка", "тип")):
                            for ci, v in enumerate(row or []):
                                vstr = str(v or "").lower()
                                if "марка" in vstr or "тип" in vstr or "таг" in vstr:
                                    tag_col = ci if tag_col < 0 else tag_col
                                if "артикул" in vstr or "код" in vstr:
                                    art_col = ci if art_col < 0 else art_col
                                if "наименование" in vstr or "изделие" in vstr or "описание" in vstr:
                                    name_col = ci
                                if "кол" in vstr or "шт" in vstr or "ед" in vstr or "количество" in vstr:
                                    qty_col = ci
                            if name_col >= 0 or art_col >= 0 or tag_col >= 0:
                                data_start = ri + 1
                                break

            for ri in range(data_start, len(rows_list)):
                row = rows_list[ri]
                row_str = " ".join(str(v or "").lower() for v in (row or []))
                if "итого" in row_str or "total" in row_str:
                    break
                item_no = _cell(row, 1)
                tag = _cell(row, tag_col) if 0 <= tag_col < len(row) else ""
                art = _cell(row, art_col) if 0 <= art_col < len(row) else ""
                package_no = _cell(row, package_col) if 0 <= package_col < len(row) else ""
                name = _cell(row, name_col) if 0 <= name_col < len(row) else ""
                qty_s = _cell(row, qty_col) if 0 <= qty_col < len(row) else "1"
                mfr = _cell(row, mfr_col) if 0 <= mfr_col < len(row) else ""
                country = _cell(row, country_col) if 0 <= country_col < len(row) else ""
                if not (name or tag or art):
                    continue
                name_lower = name.lower() if name else ""
                if name_lower and any(
                    h in name_lower for h in (
                        "наименование товара:", "goods description:", "частичная отправка:", "partial shipment:",
                        "номер грузового места", "package number", "размеры ", "dimension ",
                        "вес нетто", "net weight", "вес брутто", "gross weight", "подпись", "signature",
                        "отгрузка", "shipment", "грузового места из", "из упаковочного",
                    )
                ) and not (tag or art):
                    continue
                try:
                    qty = float(str(qty_s).replace(",", ".")) if qty_s else 1.0
                except (ValueError, TypeError):
                    qty = 1.0
                nomenclature_items.append({
                    "tag_number": tag.replace("\n", " ").strip(),
                    "package_number": package_no.replace("\n", " ").strip() if package_no else "",
                    "article": str(art).replace("\n", " ").strip() if art else "",
                    "name": name.replace("\n", " ").strip() if name else "",
                    "quantity": qty,
                    "manufacturer": mfr.replace("\n", " ").strip() if mfr else "",
                    "country_of_origin": country.replace("\n", " ").strip() if country else "",
                    "row_index": ri,
                })

        wb.close()
        header = {k: v for k, v in header.items() if v}
        return {"header": header, "nomenclature_items": nomenclature_items} if (header or nomenclature_items) else None
    except Exception:
        return None


def _extract_date_from_text(text: str) -> date | None:
    """Извлечь дату из текста (DD.MM.YYYY, DD/MM/YYYY, YYYY-MM-DD)."""
    if not text:
        return None
    patterns = [
        r"\b(\d{1,2})\.(\d{1,2})\.(\d{4})\b",
        r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b",
        r"\b(\d{4})-(\d{2})-(\d{2})\b",
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            g = m.groups()
            try:
                if len(g[0]) == 4:  # YYYY-MM-DD
                    return date(int(g[0]), int(g[1]), int(g[2]))
                return date(int(g[2]), int(g[1]), int(g[0]))
            except (ValueError, IndexError):
                continue
    return None


def _generate_supply_template_xlsx() -> bytes:
    """Сгенерировать шаблон XLSX для отгрузочной документации (ведомость грузового места)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    wb = Workbook()
    ws = wb.active
    ws.title = "Ведомость"
    # Заголовок документа
    ws.append(["Ведомость грузового места / Packing note", "", "", "", "", ""])
    ws.append(["Договор / Contract", "", "", "", "", ""])
    ws.append(["Продавец / Seller", "", "", "", "Покупатель / Buyer", ""])
    ws.append(["Адрес отгрузки / Address of pickup", "", "", "", "Адрес поставки / Delivery address", ""])
    ws.append(["Дата отгрузки", "", "", "", "", ""])
    ws.append([])
    # Заголовки таблицы номенклатуры
    headers = ["№", "Таговый номер / Tag number", "Артикул / Vendor code", "Наименование товара / Product name", "Кол-во / Qty", "Производитель / Manufacturer", "№ груз. места"]
    ws.append(headers)
    # Примеры строк
    ws.append([1, "ТН-001", "МП-100", "Манометр показывающий 0-1 МПа", 2, "Завод Манометр", "GM-001"])
    ws.append([2, "ТН-002", "МП-160", "Манометр показывающий 0-2.5 МПа", 1, "Завод Манометр", "GM-001"])
    ws.append([3, "", "", "", "", "", ""])
    for col in range(1, len(headers) + 1):
        ws.cell(1, col).font = Font(bold=True)
        ws.cell(9, col).font = Font(bold=True)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


@router.get("/template/supplies")
async def download_supply_template(
    user: User = Depends(require_trial_or_subscription),
):
    """Скачать шаблон XLSX для отгрузочной документации."""
    content = _generate_supply_template_xlsx()
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="shablon_otgruzochnaya_vedomost.xlsx"'},
    )


def _significant_words(t: str) -> list[str]:
    """Извлечь значимые слова (≥3 символов) из текста для обратного поиска."""
    return [w.lower() for w in re.findall(r"[а-яёa-z0-9\-]+", (t or "").lower()) if len(w) >= 3]


def _compute_match_score(search_term: str, text: str, is_code: bool = False) -> int:
    """Оценка совпадения 0-100. search_term и text в нижнем регистре для поиска."""
    if not search_term or not text:
        return 0
    st = search_term.lower().strip()
    tx = text.lower()
    if not st or len(st) < 2:
        return 0
    if st in tx:
        if is_code:
            return 95 if len(st) >= 4 else 85
        return 80
    words = st.split()
    if len(words) > 1:
        if all(w in tx for w in words if len(w) >= 2):
            return 70
    if len(st) >= 4 and st[:4] in tx:
        return 50
    return 0


@router.post("/recognize")
async def recognize_documents(
    files: Annotated[list[UploadFile], File(description="Файлы для распознавания (PDF, DOCX, XLSX)")],
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Распознать документы: извлечь текст и предложить привязку к сущностям по кодам/названиям."""
    try:
        return await _recognize_documents_impl(files, user, db)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("recognize 500: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


async def _recognize_documents_impl(files, user, db):
    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        raise HTTPException(403, "Нет доступа к компании")
    filter_ids = company_ids

    nom_q = select(Nomenclature).where(Nomenclature.is_deleted == False, Nomenclature.company_id.in_(filter_ids))
    supp_q = select(Supplier).where(Supplier.is_deleted == False, Supplier.company_id.in_(filter_ids))
    mfr_q = select(Manufacturer).where(Manufacturer.is_deleted == False, Manufacturer.company_id.in_(filter_ids))
    supply_q = select(Supply).where(Supply.is_deleted == False, Supply.company_id.in_(filter_ids))
    contract_q = select(Contract).where(Contract.is_deleted == False, Contract.company_id.in_(filter_ids))

    noms = (await db.execute(nom_q)).scalars().all()
    suppliers = (await db.execute(supp_q)).scalars().all()
    manufacturers = (await db.execute(mfr_q)).scalars().all()
    supplies = (await db.execute(supply_q)).scalars().all()
    contracts = (await db.execute(contract_q)).scalars().all()

    flat_items: list[tuple[str, bytes, str]] = []
    for uf in files:
        content = await uf.read()
        mime = uf.content_type or "application/octet-stream"
        ext = (Path(uf.filename or "").suffix or "").lower()
        if mime == "application/zip" or ext == ".zip":
            extracted = _extract_files_from_zip(content, uf.filename or "archive.zip")
            flat_items.extend(extracted)
        else:
            if _is_allowed_file(mime, ext):
                flat_items.append((uf.filename or "document", content, mime))

    ENTITY_LABELS = {
        "nomenclature_code": "Номенклатура (код)",
        "nomenclature_name": "Номенклатура (наименование)",
        "supplier": "Поставщик",
        "manufacturer": "Производитель",
        "contract": "Договор",
        "supply": "Поставка",
    }

    result_list = []
    for filename, content, mime in flat_items:
        ext = Path(filename or "").suffix.lower()
        text = _extract_text_from_file(content, filename, mime)
        # Добавляем имя файла (без расширения) — часто содержит название номенклатуры
        stem = Path(filename or "").stem or ""
        if stem:
            text = (text or "") + " " + stem.replace("_", " ").replace("-", " ")
        extracted_preview = (text[:500] + "...") if len(text) > 500 else text if text else ""
        suggested: list[dict] = []
        extracted_structured: dict | None = None
        detected_columns: list[dict] = []

        universal = _extract_universal(content, filename, mime)
        detected_columns = universal.get("detected_columns", [])
        rows_data = universal.get("rows_data", [])
        uni_header = universal.get("header", {})
        if detected_columns:
            extracted_structured = extracted_structured or {}
            extracted_structured["detected_columns"] = detected_columns
            extracted_structured["raw_preview"] = universal.get("raw_preview", "")
            extracted_structured["rows_count"] = len(rows_data)
            extracted_structured["rows_data"] = rows_data
            if uni_header:
                extracted_structured["header"] = {**extracted_structured.get("header", {}), **uni_header}

        if ext == ".xlsx":
            packing = _parse_packing_document_xlsx(content, filename)
            nom_items_raw = packing.get("nomenclature_items", []) if packing else []
            if packing and packing.get("header"):
                extracted_structured = extracted_structured or {}
                extracted_structured["header"] = packing.get("header", {})
            if nom_items_raw or (rows_data and detected_columns):
                extracted_structured = extracted_structured or {}
                if not nom_items_raw and rows_data and detected_columns:
                    nom_items_raw = _build_nom_items_from_universal(rows_data, detected_columns)
                if nom_items_raw:
                    header = extracted_structured.get("header", {})
                    header_text = " ".join(str(v) for v in header.values())
                    supply_date_detected = _extract_date_from_text(header_text)
                    extracted_structured["supply_document"] = True
                    extracted_structured["supply_date_detected"] = supply_date_detected.isoformat() if supply_date_detected else None
                nom_items_with_matches = []
                for it in nom_items_raw:
                    item_copy = dict(it)
                    code_val = (it.get("article") or it.get("code") or "").strip()
                    tag_val = (it.get("tag_number") or "").strip()
                    name_val = (it.get("name") or "").strip()
                    match_candidates = []
                    conflict_same_code_diff_tag = False
                    for n in noms:
                        score = 0
                        match_type = ""
                        if code_val and n.code and code_val.lower() == (n.code or "").lower():
                            if tag_val and n.tag_number and tag_val.lower() == (n.tag_number or "").lower():
                                score = 100
                                match_type = "exact"
                            elif tag_val or (n.tag_number and n.tag_number.strip()):
                                if (tag_val or "").lower() != (n.tag_number or "").lower():
                                    conflict_same_code_diff_tag = True
                                score = 85
                                match_type = "code_match"
                        elif code_val and n.code and _compute_match_score(code_val, n.code or "", is_code=True) >= 50:
                            score = max(score, 70)
                            if not match_type:
                                match_type = "code_fuzzy"
                        elif name_val and n.name and _compute_match_score(name_val, n.name or "", is_code=False) >= 50:
                            score = max(score, _compute_match_score(name_val, n.name or "", is_code=False))
                            if not match_type:
                                match_type = "name"
                        if score >= 50:
                            match_candidates.append({
                                "nomenclature_id": n.id,
                                "name": n.name,
                                "code": n.code or "",
                                "tag_number": n.tag_number or "",
                                "match_score": min(100, score + 5),
                                "match_type": match_type,
                            })
                    match_candidates.sort(key=lambda x: -x["match_score"])
                    item_copy["match_candidates"] = match_candidates[:5]
                    item_copy["conflict_same_code_diff_tag"] = conflict_same_code_diff_tag
                    nom_items_with_matches.append(item_copy)
                extracted_structured["nomenclature_items"] = nom_items_with_matches
                if not detected_columns and nom_items_raw:
                    for it in nom_items_raw[:1]:
                        for k, v in it.items():
                            if v:
                                detected_columns.append({"index": len(detected_columns), "name": k, "example": str(v)[:200], "samples": []})
                header = extracted_structured.get("header", {})
                nom_items = extracted_structured.get("nomenclature_items", [])
                text = text + " " + " ".join(str(v) for v in header.values())
                for it in nom_items:
                    text = text + " " + (it.get("tag_number") or "") + " " + (it.get("article") or "") + " " + (it.get("name") or "") + " " + (it.get("manufacturer") or "")

        text_words = _significant_words(text)

        for n in noms:
            score_code = _compute_match_score(n.code, text, is_code=True) if n.code and n.code.strip() else 0
            score_name = _compute_match_score(n.name, text, is_code=False) if n.name and len(n.name) >= 3 else 0
            score_rev = 0
            for w in text_words:
                if n.name and w in (n.name or "").lower():
                    score_rev = max(score_rev, 75)
                if n.code and w in (n.code or "").lower():
                    score_rev = max(score_rev, 80)
            score = max(score_code, score_name, score_rev)
            if score >= 50:
                if score_code >= score_name and n.code:
                    field_label = ENTITY_LABELS["nomenclature_code"]
                    matched_term = n.code
                else:
                    field_label = ENTITY_LABELS["nomenclature_name"]
                    matched_term = n.name or ""
                suggested.append({
                    "entity_type": "nomenclature",
                    "entity_id": n.id,
                    "name": n.name,
                    "code": n.code or "",
                    "match_score": min(100, score + 5),
                    "field_label": field_label,
                    "matched_term": matched_term,
                })

        for s in suppliers:
            if s.name and len(s.name) >= 3:
                score = _compute_match_score(s.name, text, is_code=False)
                if score >= 50:
                    suggested.append({
                        "entity_type": "supplier",
                        "entity_id": s.id,
                        "name": s.name,
                        "code": s.inn or "",
                        "match_score": min(100, score + 5),
                        "field_label": ENTITY_LABELS["supplier"],
                        "matched_term": s.name,
                    })

        for m in manufacturers:
            if m.name and len(m.name) >= 3:
                score = _compute_match_score(m.name, text, is_code=False)
                if score >= 50:
                    suggested.append({
                        "entity_type": "manufacturer",
                        "entity_id": m.id,
                        "name": m.name,
                        "code": "",
                        "match_score": min(100, score + 5),
                        "field_label": ENTITY_LABELS["manufacturer"],
                        "matched_term": m.name,
                    })

        for c in contracts:
            if c.number and c.number.strip():
                score = _compute_match_score(c.number, text, is_code=True)
                if score >= 50:
                    suggested.append({
                        "entity_type": "contract",
                        "entity_id": c.id,
                        "name": c.number or f"Договор #{c.id}",
                        "code": c.number or "",
                        "match_score": min(100, score + 5),
                        "field_label": ENTITY_LABELS["contract"],
                        "matched_term": c.number or "",
                    })

        suggested.sort(key=lambda x: -x["match_score"])
        suggested = suggested[:15]
        out: dict = {
            "filename": filename,
            "extracted_preview": extracted_preview,
            "suggested": suggested,
        }
        if extracted_structured:
            out["extracted_structured"] = extracted_structured
        result_list.append(out)

    return {"files": result_list}


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    entity_type: str = Form(...),
    entity_id: int = Form(...),
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Загрузить документ к сущности."""
    if entity_type not in ENTITY_TYPES:
        raise HTTPException(400, f"Недопустимый entity_type: {entity_type}")
    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        raise HTTPException(403, "Нет доступа к компании")
    company_id = company_ids[0]
    content = await file.read()
    _check_file_size(content)
    mime = file.content_type or "application/octet-stream"
    ext = (Path(file.filename or "").suffix or "").lower()
    _verify_magic_bytes(content, mime, ext)
    if not _is_allowed_file(mime, ext):
        raise HTTPException(400, f"Недопустимый тип файла: {mime}. Разрешены: PDF, изображения, Word, Excel, ZIP.")
    result = await _save_document(
        content, file.filename or "doc", mime, entity_type, entity_id, company_id, user.id, db
    )
    await db.commit()
    return result


@router.post("/upload-batch")
async def upload_batch(
    files: Annotated[list[UploadFile], File(description="Файлы для загрузки (ZIP будет распакован)")],
    entity_type: Annotated[str | None, Form()] = None,
    entity_id: Annotated[int | None, Form()] = None,
    assignments: Annotated[str | None, Form(description="JSON: [{\"file_index\": 0, \"entity_type\": \"nomenclature\", \"entity_id\": 1}]")] = None,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Загрузить несколько документов. Либо entity_type+entity_id для всех, либо assignments для каждого файла."""
    if entity_type is not None and entity_id is not None:
        if entity_type not in ENTITY_TYPES:
            raise HTTPException(400, f"Недопустимый entity_type: {entity_type}")
    elif assignments:
        try:
            assigns_list = json.loads(assignments)
            for a in assigns_list:
                if a.get("entity_type") not in ENTITY_TYPES:
                    raise HTTPException(400, f"Недопустимый entity_type: {a.get('entity_type')}")
        except json.JSONDecodeError as e:
            raise HTTPException(400, f"Неверный JSON в assignments: {e}")
    else:
        raise HTTPException(400, "Укажите entity_type и entity_id либо assignments (JSON)")

    company_ids = await get_user_company_ids(user, db)
    if not company_ids:
        raise HTTPException(403, "Нет доступа к компании")
    company_id = company_ids[0]

    # Развернуть файлы: обычные — как есть, ZIP — извлечь содержимое
    flat_items: list[tuple[str, bytes, str]] = []
    for uf in files:
        content = await uf.read()
        _check_file_size(content)
        mime = uf.content_type or "application/octet-stream"
        ext = (Path(uf.filename or "").suffix or "").lower()
        _verify_magic_bytes(content, mime, ext)
        if mime == "application/zip" or ext == ".zip":
            extracted = _extract_files_from_zip(content, uf.filename or "archive.zip")
            flat_items.extend(extracted)
        else:
            if not _is_allowed_file(mime, ext):
                continue
            flat_items.append((uf.filename or "document", content, mime))

    if not flat_items:
        raise HTTPException(400, "Нет допустимых файлов для загрузки")

    assigns_list: list[tuple[int, str, int]] = []
    if assignments:
        raw = json.loads(assignments)
        for a in raw:
            idx = a.get("file_index")
            if idx is not None and isinstance(idx, int):
                et = a.get("entity_type", "nomenclature")
                eid = a.get("entity_id", 0)
                if et in ENTITY_TYPES and eid:
                    assigns_list.append((idx, et, eid))

    results = []
    for i, (filename, content, mime) in enumerate(flat_items):
        ext = (Path(filename or "").suffix or "").lower()
        try:
            _verify_magic_bytes(content, mime, ext)
        except HTTPException:
            continue
        if assigns_list:
            for (idx, et, eid) in assigns_list:
                if idx == i:
                    try:
                        r = await _save_document(content, filename, mime, et, eid, company_id, user.id, db)
                        results.append(r)
                    except Exception:
                        pass
        else:
            et, eid = entity_type, entity_id
            try:
                r = await _save_document(content, filename, mime, et, eid, company_id, user.id, db)
                results.append(r)
            except Exception:
                pass
    await db.commit()
    return {"uploaded": len(results), "items": results}


# Стандартные поля, сопоставляемые с extra_fields при импорте номенклатуры
IMPORT_EXTRA_FIELD_KEYS = {
    "shipping_number", "storage_conditions", "packaging_type", "length_cm", "width_cm", "height_cm",
    "net_weight_unit", "total_net_weight", "total_gross_weight", "price_without_vat", "price_with_vat",
}

# Маппинг newName (при «Создать новое свойство») на стандартные ключи. Подстроки — в нижнем регистре.
# Порядок важен: более длинные/специфичные первыми.
NEWNAME_TO_STANDARD: list[tuple[str, str]] = [
    ("наименование отгружаемой позиции", "name"),
    ("наименование", "name"),
    ("название", "name"),
    ("отгружаемой позиции", "name"),
    ("номер отгружаемого", "shipping_number"),
    ("номер отгружаемо", "shipping_number"),
    ("условия хранения", "storage_conditions"),
    ("тип упаковки", "packaging_type"),
    ("длина, см", "length_cm"),
    ("длина", "length_cm"),
    ("ширина, см", "width_cm"),
    ("ширина", "width_cm"),
    ("высота, см", "height_cm"),
    ("высота", "height_cm"),
    ("вес нетто за ед", "net_weight_unit"),
    ("итоговый вес нетто", "total_net_weight"),
    ("вес нетто итого", "total_net_weight"),
    ("итоговый вес брутто", "total_gross_weight"),
    ("вес брутто итого", "total_gross_weight"),
    ("цена без ндс", "price_without_vat"),
    ("цена с ндс", "price_with_vat"),
    ("количество", "quantity"),
]


async def _run_nomenclature_import_core(
    db: AsyncSession,
    job: ImportJob,
    company_id: int,
    triggered_by_user_id: int | None,
    files_data: list[tuple[bytes, str, str]],
    mappings_raw: dict,
    row_indices_by_file: dict[int, set[int]],
    category_id: int | None,
    subcategory_id: int | None,
    supplier_id_val: int | None,
    delivery_date_val: date | None,
) -> dict:
    """Ядро импорта номенклатуры. files_data = [(content, filename, mime), ...]."""
    system_fields = {"code", "name", "tag_number", "package_number", "manufacturer", "quantity"}
    created: list[dict] = []
    updated_count = 0
    supplies_created = 0
    limit_reached_nomenclature = False
    limit_reached_manufacturers = False
    limit_message: str | None = None

    def _norm(s: str | None) -> str:
        if not s:
            return ""
        return " ".join(str(s).strip().lower().split())

    rows_to_process: list[tuple[dict, float]] = []
    for fi, (content, filename, mime) in enumerate(files_data):
        ext = (Path(filename or "").suffix or "").lower()
        if ext not in (".xlsx", ".xls", ".docx", ".doc"):
            continue
        file_map = mappings_raw.get(str(fi), mappings_raw.get(fi, {}))
        if not file_map:
            continue
        uni = _extract_universal(content, filename or "doc", mime)
        rows_data = uni.get("rows_data", [])
        if not rows_data:
            continue
        allowed_indices = row_indices_by_file.get(fi)
        for ri, row in enumerate(rows_data):
            if allowed_indices is not None and ri not in allowed_indices:
                continue
            nom_data: dict = {"name": "", "extra_fields": {}}
            for col_idx_str, mapping in file_map.items():
                col_idx = int(col_idx_str) if isinstance(col_idx_str, str) else col_idx_str
                m = mapping or {}
                map_to = m.get("mapTo") or m.get("mapto", "")
                if not map_to:
                    continue
                val = row.get(col_idx) or row.get(str(col_idx)) if isinstance(row, dict) else ""
                val = val if val is not None else ""
                if not val or not str(val).strip():
                    continue
                val = str(val).strip()
                if map_to in system_fields and map_to != "manufacturer":
                    if map_to == "quantity":
                        try:
                            nom_data.setdefault("extra_fields", {})["quantity"] = float(str(val).replace(",", "."))
                        except (ValueError, TypeError):
                            pass
                    else:
                        nom_data[map_to] = val
                elif map_to == "manufacturer":
                    nom_data["manufacturer_name"] = val
                elif map_to in IMPORT_EXTRA_FIELD_KEYS:
                    nom_data.setdefault("extra_fields", {})[map_to] = val
                elif map_to == "new":
                    new_name = (mapping or {}).get("newName", "").strip()
                    if new_name:
                        new_lower = new_name.lower()
                        mapped = None
                        for substr, std_key in NEWNAME_TO_STANDARD:
                            if substr in new_lower:
                                mapped = std_key
                                break
                        if mapped == "name":
                            if not (nom_data.get("name") or "").strip():
                                nom_data["name"] = val
                        elif mapped and mapped in IMPORT_EXTRA_FIELD_KEYS:
                            nom_data.setdefault("extra_fields", {})[mapped] = val
                        elif mapped == "quantity":
                            try:
                                nom_data.setdefault("extra_fields", {})["quantity"] = float(str(val).replace(",", "."))
                            except (ValueError, TypeError):
                                nom_data.setdefault("extra_fields", {})[new_name] = val
                        else:
                            nom_data.setdefault("extra_fields", {})[new_name] = val
                elif map_to and map_to not in ("", "— не импортировать —"):
                    map_lower = map_to.lower()
                    resolved_key = None
                    for substr, std_key in NEWNAME_TO_STANDARD:
                        if substr in map_lower:
                            resolved_key = std_key
                            break
                    if resolved_key == "name":
                        if not (nom_data.get("name") or "").strip():
                            nom_data["name"] = val
                    elif resolved_key and resolved_key in IMPORT_EXTRA_FIELD_KEYS:
                        nom_data.setdefault("extra_fields", {})[resolved_key] = val
                    elif resolved_key == "quantity":
                        try:
                            nom_data.setdefault("extra_fields", {})["quantity"] = float(str(val).replace(",", "."))
                        except (ValueError, TypeError):
                            nom_data.setdefault("extra_fields", {})[map_to] = val
                    else:
                        nom_data.setdefault("extra_fields", {})[map_to] = val

            name = (nom_data.get("name") or nom_data.get("tag_number") or nom_data.get("code") or "").strip()
            if not name:
                continue
            qty = 1.0
            try:
                qty = float(nom_data.get("extra_fields", {}).get("quantity", 1))
            except (TypeError, ValueError):
                pass
            rows_to_process.append((nom_data, qty))

    supply_items: list[dict] = []
    total_supply_qty = 0.0
    for nom_data, row_qty in rows_to_process:
        code_val = (nom_data.get("code") or "").strip() or None
        tag_val = (nom_data.get("tag_number") or "").strip() or None
        name_val = (nom_data.get("name") or nom_data.get("tag_number") or nom_data.get("code") or "").strip()[:255] or None
        manufacturer_id_val: int | None = None
        mfr_name = nom_data.get("manufacturer_name", "").strip()
        if mfr_name:
            mfr_norm = mfr_name[:255].strip()
            mfr_q = select(Manufacturer).where(
                Manufacturer.company_id == company_id,
                Manufacturer.is_deleted == False,
                Manufacturer.name == mfr_norm,
            )
            mfr_r = await db.execute(mfr_q)
            mfr_existing = mfr_r.scalar_one_or_none()
            if mfr_existing:
                manufacturer_id_val = mfr_existing.id
            else:
                try:
                    await check_entity_limit(company_id, "manufacturers", db)
                except HTTPException as e:
                    if e.status_code == 403:
                        limit_reached_manufacturers = True
                        manufacturer_id_val = None
                    else:
                        raise
                else:
                    mfr_new = Manufacturer(company_id=company_id, name=mfr_name[:255])
                    db.add(mfr_new)
                    await db.flush()
                    manufacturer_id_val = mfr_new.id

        extra = nom_data.get("extra_fields") or {}
        if "manufacturer" in extra:
            del extra["manufacturer"]

        existing_nom = None
        if code_val or tag_val:
            q = select(Nomenclature).where(
                Nomenclature.company_id == company_id,
                Nomenclature.is_deleted == False,
            )
            if code_val and tag_val:
                q = q.where(Nomenclature.code == code_val, Nomenclature.tag_number == tag_val)
            elif code_val:
                q = q.where(Nomenclature.code == code_val)
            else:
                q = q.where(Nomenclature.tag_number == tag_val)
            r = await db.execute(q)
            existing_nom = r.scalars().first()
        if not existing_nom and name_val:
            name_norm = _norm(name_val)
            all_noms = (await db.execute(
                select(Nomenclature).where(
                    Nomenclature.company_id == company_id,
                    Nomenclature.is_deleted == False,
                )
            ))
            for n in all_noms.scalars().all():
                if _norm(n.name) == name_norm or (n.name and n.name.strip() == name_val):
                    existing_nom = n
                    break

        if existing_nom:
            merged = {**(existing_nom.extra_fields or {}), **extra}
            existing_nom.extra_fields = merged if merged else None
            if category_id is not None:
                existing_nom.category_id = category_id
            if subcategory_id is not None:
                existing_nom.subcategory_id = subcategory_id
            if manufacturer_id_val is not None:
                existing_nom.manufacturer_id = manufacturer_id_val
            if (nom_data.get("package_number") or "").strip():
                existing_nom.package_number = (nom_data.get("package_number") or "").strip()[:100] or None
            nom = existing_nom
            updated_count += 1
        else:
            try:
                await check_entity_limit(company_id, "nomenclature", db)
                if supplier_id_val:
                    await check_nomenclature_per_supplier_limit(company_id, supplier_id_val, db)
            except HTTPException as e:
                if e.status_code == 403:
                    limit_reached_nomenclature = True
                    limit_message = e.detail or "Достигнут лимит. Оформите подписку для расширения."
                    continue
                raise
            nom = Nomenclature(
                company_id=company_id,
                code=code_val or None,
                name=name_val or "—",
                category_id=category_id,
                subcategory_id=subcategory_id,
                tag_number=tag_val or None,
                package_number=nom_data.get("package_number") or None,
                manufacturer_id=manufacturer_id_val,
                extra_fields=extra if extra else None,
            )
            db.add(nom)
            await db.flush()
            created.append({"id": nom.id, "name": nom.name, "code": nom.code})

        if supplier_id_val:
            supply_items.append({
                "nomenclature_id": nom.id,
                "quantity": row_qty,
                "name": (nom.name or nom.code or "")[:255],
            })
            total_supply_qty += row_qty

    if supplier_id_val and supply_items:
        first = supply_items[0]
        supply_extra: dict = {"items": supply_items}
        supply = Supply(
            company_id=company_id,
            supplier_id=supplier_id_val,
            nomenclature_id=first["nomenclature_id"],
            quantity=total_supply_qty,
            delivery_date=delivery_date_val,
            extra_fields=supply_extra,
        )
        db.add(supply)
        supplies_created = 1

    job.status = "completed"
    job.finished_at = datetime.utcnow()
    job.stats = {
        "created": len(created),
        "updated": updated_count,
        "supplies": supplies_created,
        "limit_reached": limit_reached_nomenclature or limit_reached_manufacturers,
    }
    return {
        "created": len(created),
        "updated": updated_count,
        "supplies": supplies_created,
        "items": created,
        "limit_reached": limit_reached_nomenclature or limit_reached_manufacturers,
        "limit_message": limit_message if (limit_reached_nomenclature or limit_reached_manufacturers) else None,
    }


@router.post("/import-nomenclature")
async def import_nomenclature(
    request: Request,
    files: Annotated[list[UploadFile], File(description="Файлы Excel/Word для импорта")],
    column_mappings: Annotated[str, Form(description="JSON: {fileIndex: {colIndex: {mapTo, newName?}}}")],
    category_id: Annotated[int | None, Form()] = None,
    subcategory_id: Annotated[int | None, Form()] = None,
    supplier_id: Annotated[str | None, Form()] = None,
    delivery_date: Annotated[str | None, Form()] = None,
    row_indices: Annotated[str | None, Form()] = None,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Импорт номенклатуры из файлов по сопоставлению колонок. При supplier_id создаются поставки."""
    company_ids = await get_user_company_ids(user, db)
    if company_ids is not None and not company_ids:
        raise HTTPException(403, "Нет доступа к компании")
    company_id = company_ids[0] if company_ids else None
    if not company_id:
        raise HTTPException(403, "Нет доступа к компании")
    form = await request.form()
    _supplier_id = supplier_id
    if (not _supplier_id or not str(_supplier_id).strip()) and "supplier_id" in form:
        v = form.get("supplier_id")
        if v is not None and not hasattr(v, "filename"):
            _supplier_id = str(v).strip() if v else None
    _delivery_date = delivery_date
    if (not _delivery_date or not str(_delivery_date).strip()) and "delivery_date" in form:
        v = form.get("delivery_date")
        if v is not None and not hasattr(v, "filename"):
            _delivery_date = str(v).strip() if v else None
    try:
        mappings_raw = json.loads(column_mappings)
    except json.JSONDecodeError as e:
        raise HTTPException(400, f"Неверный JSON в column_mappings: {e}")

    row_indices_by_file: dict[int, set[int]] = {}
    if row_indices:
        try:
            raw = json.loads(row_indices)
            for k, v in raw.items():
                fi = int(k) if isinstance(k, str) else k
                row_indices_by_file[fi] = set(int(i) for i in v) if isinstance(v, (list, tuple)) else {int(v)}
        except (json.JSONDecodeError, ValueError):
            pass

    delivery_date_val: date | None = None
    if _delivery_date and str(_delivery_date).strip():
        try:
            from datetime import datetime as dt
            s = str(_delivery_date).strip()[:10]
            if len(s) >= 10 and s[2] in ("-", ".") and s[5] in ("-", "."):
                s = s.replace(".", "-")
            delivery_date_val = dt.strptime(s[:10], "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    supplier_id_val: int | None = None
    if _supplier_id and str(_supplier_id).strip():
        try:
            sid = int(str(_supplier_id).strip())
            q = select(Supplier).where(Supplier.id == sid, Supplier.is_deleted == False)
            if company_ids is not None:
                q = q.where(Supplier.company_id.in_(company_ids))
            r = await db.execute(q)
            if r.scalar_one_or_none():
                supplier_id_val = sid
        except (ValueError, TypeError):
            pass
    sub_url = _cabinet_payment_subscription_url()

    job = ImportJob(
        company_id=company_id,
        job_type="nomenclature_import",
        status="running",
        triggered_by_user_id=user.id,
    )
    db.add(job)
    await db.flush()

    files_data: list[tuple[bytes, str, str]] = []
    for uf in files:
        content = await uf.read()
        mime = uf.content_type or "application/octet-stream"
        files_data.append((content, uf.filename or "doc", mime))

    result = await _run_nomenclature_import_core(
        db, job, company_id, user.id,
        files_data, mappings_raw, row_indices_by_file,
        category_id, subcategory_id, supplier_id_val, delivery_date_val,
    )
    await db.commit()

    limit_reached = result.get("limit_reached", False)
    limit_message = result.get("limit_message")
    created = result.get("items", [])
    msg = f"Создано позиций: {len(created)}, обновлено: {result.get('updated', 0)}, поставок: {result.get('supplies', 0)}"
    if limit_reached and limit_message:
        msg += f". {limit_message}"
    return {
        "created": result.get("created", 0),
        "updated": result.get("updated", 0),
        "supplies": result.get("supplies", 0),
        "items": created,
        "message": msg,
        "limit_reached": limit_reached,
        "limit_message": limit_message if limit_reached else None,
        "subscription_url": sub_url if limit_reached else None,
    }


@router.post("/import-supplies")
async def import_supplies(
    files: Annotated[list[UploadFile], File(description="Файлы отгрузочной документации (XLSX)")],
    supplier_id: Annotated[int | None, Form(description="ID поставщика")] = None,
    use_supply_date: Annotated[bool, Form(description="Использовать дату поставки")] = False,
    supply_date: Annotated[str | None, Form(description="Дата поставки YYYY-MM-DD (если use_supply_date=True)")] = None,
    contract_id: Annotated[int | None, Form(description="ID договора")] = None,
    buyer_name: Annotated[str | None, Form(description="Покупатель (текст)")] = None,
    shipping_address: Annotated[str | None, Form(description="Адрес отгрузки")] = None,
    delivery_address: Annotated[str | None, Form(description="Адрес доставки")] = None,
    items: Annotated[str, Form(description="JSON: список {row_index, action: 'merge'|'new', nomenclature_id?, quantity, code?, name?, tag_number?, package_number?, manufacturer_name?}")] = "[]",
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Импорт поставок из файла с распознаванием. action=merge — добавить к существующей номенклатуре, new — создать новую и поставить."""
    company_id = await get_user_company_id(user, db)
    if not company_id:
        raise HTTPException(403, "Нет доступа к компании")
    if not supplier_id and items != "[]":
        raise HTTPException(400, "Укажите поставщика")
    try:
        items_list = json.loads(items)
    except json.JSONDecodeError as e:
        raise HTTPException(400, f"Неверный JSON в items: {e}")
    prod_date = None
    delivery_date_val: date | None = None
    if use_supply_date and supply_date:
        try:
            from datetime import datetime as dt
            s = str(supply_date).strip()[:10]
            if len(s) >= 10:
                if s[4] == "-" and s[7] == "-":  # yyyy-mm-dd
                    delivery_date_val = date.fromisoformat(s[:10])
                elif s[2] in (".", "-") and s[5] in (".", "-"):  # dd.mm.yyyy
                    delivery_date_val = dt.strptime(s[:10].replace(".", "-"), "%d-%m-%Y").date()
                else:
                    delivery_date_val = date.fromisoformat(s[:10])
                prod_date = delivery_date_val  # для совместимости
        except (ValueError, TypeError):
            pass
    supply_extra: dict = {}
    if contract_id:
        supply_extra["contract_id"] = contract_id
    if buyer_name and buyer_name.strip():
        supply_extra["buyer_name"] = buyer_name.strip()[:255]
    if shipping_address and shipping_address.strip():
        supply_extra["shipping_address"] = shipping_address.strip()[:512]
    if delivery_address and delivery_address.strip():
        supply_extra["delivery_address"] = delivery_address.strip()[:512]
    created_supplies: list[dict] = []
    limit_reached_nomenclature = False
    limit_reached_manufacturers = False
    limit_message: str | None = None
    sub_url = _cabinet_payment_subscription_url()
    for uf in files:
        content = await uf.read()
        mime = uf.content_type or "application/octet-stream"
        ext = (Path(uf.filename or "").suffix or "").lower()
        if ext not in (".xlsx", ".xls"):
            continue
        packing = _parse_packing_document_xlsx(content, uf.filename or "doc.xlsx")
        nom_items = packing.get("nomenclature_items", []) if packing else []
        if not nom_items:
            universal = _extract_universal(content, uf.filename or "doc.xlsx", mime)
            rows_data = universal.get("rows_data", [])
            detected_columns = universal.get("detected_columns", [])
            if rows_data and detected_columns:
                nom_items = _build_nom_items_from_universal(rows_data, detected_columns)
        if not nom_items:
            continue
        # Собираем все позиции в одну поставку на файл
        supply_items: list[dict] = []
        total_qty = 0.0
        # merge: агрегируем по nomenclature_id
        merged_qty: dict[int, float] = {}
        for itm in items_list:
            row_idx = itm.get("row_index", -1)
            action = itm.get("action", "new")
            nom_id = itm.get("nomenclature_id")
            qty = float(itm.get("quantity", 1))
            if row_idx < 0 or row_idx >= len(nom_items):
                continue
            if action == "merge" and nom_id:
                merged_qty[nom_id] = merged_qty.get(nom_id, 0) + qty
        for nom_id, q in merged_qty.items():
            nom_r = await db.execute(select(Nomenclature.name, Nomenclature.code).where(Nomenclature.id == nom_id))
            nom_row = nom_r.one_or_none()
            nm = (nom_row.name or nom_row.code or "")[:255] if nom_row else ""
            supply_items.append({"nomenclature_id": nom_id, "quantity": q, "name": nm})
            total_qty += q
        # new: создаём номенклатуру и добавляем в позиции
        for itm in items_list:
            row_idx = itm.get("row_index", -1)
            action = itm.get("action", "new")
            nom_id = itm.get("nomenclature_id")
            qty = float(itm.get("quantity", 1))
            if row_idx < 0 or row_idx >= len(nom_items):
                continue
            src = nom_items[row_idx]
            if action == "merge" and nom_id:
                continue
            elif action == "new":
                code_val = (itm.get("code") or src.get("article") or "").strip() or None
                name_val = (itm.get("name") or src.get("name") or src.get("tag_number") or "").strip()
                if not name_val:
                    continue
                if limit_reached_nomenclature:
                    continue
                try:
                    await check_entity_limit(company_id, "nomenclature", db)
                    if supplier_id:
                        await check_nomenclature_per_supplier_limit(company_id, supplier_id, db)
                except HTTPException as e:
                    if e.status_code == 403:
                        limit_reached_nomenclature = True
                        limit_message = e.detail or "Достигнут лимит. Оформите подписку для расширения."
                        continue
                    raise
                tag_val = (itm.get("tag_number") or src.get("tag_number") or "").strip() or None
                pkg_val = (itm.get("package_number") or src.get("package_number") or "").strip() or None
                mfr_name = (itm.get("manufacturer_name") or src.get("manufacturer") or "").strip()
                manufacturer_id_val = None
                if mfr_name:
                    mfr_q = select(Manufacturer).where(
                        Manufacturer.company_id == company_id,
                        Manufacturer.is_deleted == False,
                        Manufacturer.name == mfr_name[:255],
                    )
                    mfr_r = await db.execute(mfr_q)
                    mfr_existing = mfr_r.scalar_one_or_none()
                    if mfr_existing:
                        manufacturer_id_val = mfr_existing.id
                    else:
                        try:
                            await check_entity_limit(company_id, "manufacturers", db)
                        except HTTPException as e:
                            if e.status_code == 403:
                                limit_reached_manufacturers = True
                                if not limit_message:
                                    limit_message = e.detail or "Достигнут лимит производителей."
                            else:
                                raise
                        else:
                            mfr_new = Manufacturer(company_id=company_id, name=mfr_name[:255])
                            db.add(mfr_new)
                            await db.flush()
                            manufacturer_id_val = mfr_new.id
                nom = Nomenclature(
                    company_id=company_id,
                    code=code_val,
                    name=name_val[:255],
                    tag_number=tag_val,
                    package_number=pkg_val,
                    manufacturer_id=manufacturer_id_val,
                )
                db.add(nom)
                await db.flush()
                supply_items.append({"nomenclature_id": nom.id, "quantity": qty, "name": nom.name or ""})
                total_qty += qty
        if supply_items and supplier_id:
            sup_extra = {**supply_extra, "items": supply_items}
            first = supply_items[0]
            sup = Supply(
                company_id=company_id,
                supplier_id=supplier_id,
                nomenclature_id=first["nomenclature_id"],
                quantity=total_qty,
                production_date=prod_date,
                delivery_date=delivery_date_val,
                calibration_date=None,
                extra_fields=sup_extra if sup_extra else None,
            )
            db.add(sup)
            await db.flush()
            created_supplies.append({"id": sup.id, "nomenclature_id": first["nomenclature_id"], "quantity": total_qty})
    await db.commit()
    limit_reached = limit_reached_nomenclature or limit_reached_manufacturers
    return {
        "created": len(created_supplies),
        "supplies": created_supplies,
        "limit_reached": limit_reached,
        "limit_message": limit_message if limit_reached else None,
        "subscription_url": sub_url if limit_reached else None,
    }


@router.get("")
async def list_documents(
    entity_type: str | None = None,
    entity_id: int | None = None,
    search: str | None = Query(None, description="Поиск по названию файла"),
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Список документов по сущности, компании или поиску. При фильтре по производителю/категории
    учитываются документы связанной номенклатуры (документы загружаются в карточках номенклатуры)."""
    company_ids = await get_user_company_ids(user, db)
    q = select(Document)
    if company_ids is not None:
        q = q.where(Document.company_id.in_(company_ids))

    if entity_type and entity_id is not None:
        eid = entity_id
        if entity_type == "manufacturer":
            # Документы производителя напрямую ИЛИ документы номенклатуры от этого производителя
            nom_ids_subq = select(Nomenclature.id).where(
                Nomenclature.manufacturer_id == eid,
                Nomenclature.is_deleted == False,
            )
            if company_ids is not None:
                nom_ids_subq = nom_ids_subq.where(Nomenclature.company_id.in_(company_ids))
            nom_ids = [r[0] for r in (await db.execute(nom_ids_subq)).all()]
            conds = [and_(Document.entity_type == "manufacturer", Document.entity_id == eid)]
            if nom_ids:
                conds.append(and_(Document.entity_type == "nomenclature", Document.entity_id.in_(nom_ids)))
            q = q.where(or_(*conds))
        elif entity_type == "supplier":
            # Документы поставщика, его поставок и договоров; + номенклатура из его поставок
            supply_ids_subq = select(Supply.id).where(Supply.supplier_id == eid, Supply.is_deleted == False)
            if company_ids is not None:
                supply_ids_subq = supply_ids_subq.where(Supply.company_id.in_(company_ids))
            supply_ids = [r[0] for r in (await db.execute(supply_ids_subq)).all()]
            contract_ids_subq = select(Contract.id).where(
                Contract.supplier_id == eid,
                Contract.is_deleted == False,
            )
            if company_ids is not None:
                contract_ids_subq = contract_ids_subq.where(Contract.company_id.in_(company_ids))
            contract_ids = [r[0] for r in (await db.execute(contract_ids_subq)).all()]
            nom_ids_subq = select(Supply.nomenclature_id).where(
                Supply.supplier_id == eid,
                Supply.is_deleted == False,
                Supply.nomenclature_id.isnot(None),
            )
            if company_ids is not None:
                nom_ids_subq = nom_ids_subq.where(Supply.company_id.in_(company_ids))
            nom_ids = list({r[0] for r in (await db.execute(nom_ids_subq)).all() if r[0]})
            conds = [
                and_(Document.entity_type == "supplier", Document.entity_id == eid),
            ]
            if supply_ids:
                conds.append(and_(Document.entity_type == "supply", Document.entity_id.in_(supply_ids)))
            if contract_ids:
                conds.append(and_(Document.entity_type == "contract", Document.entity_id.in_(contract_ids)))
            if nom_ids:
                conds.append(and_(Document.entity_type == "nomenclature", Document.entity_id.in_(nom_ids)))
            q = q.where(or_(*conds))
        elif entity_type == "category":
            # Документы номенклатуры в этой категории или в подкатегориях категории
            from app.models.entity import SubCategory
            subcat_ids_subq = select(SubCategory.id).join(Category, SubCategory.category_id == Category.id).where(Category.id == eid)
            if company_ids is not None:
                subcat_ids_subq = subcat_ids_subq.where(Category.company_id.in_(company_ids))
            subcat_ids = [r[0] for r in (await db.execute(subcat_ids_subq)).all()]
            cat_filter = Nomenclature.category_id == eid
            subcat_filter = Nomenclature.subcategory_id.in_(subcat_ids) if subcat_ids else Nomenclature.id < 0
            nom_ids_subq = select(Nomenclature.id).where(
                or_(cat_filter, subcat_filter),
                Nomenclature.is_deleted == False,
            )
            if company_ids is not None:
                nom_ids_subq = nom_ids_subq.where(Nomenclature.company_id.in_(company_ids))
            nom_ids = [r[0] for r in (await db.execute(nom_ids_subq)).all()]
            if nom_ids:
                q = q.where(and_(Document.entity_type == "nomenclature", Document.entity_id.in_(nom_ids)))
            else:
                q = q.where(Document.entity_id < 0)
        elif entity_type == "subcategory":
            nom_ids_subq = select(Nomenclature.id).where(Nomenclature.subcategory_id == eid, Nomenclature.is_deleted == False)
            if company_ids is not None:
                nom_ids_subq = nom_ids_subq.where(Nomenclature.company_id.in_(company_ids))
            nom_ids = [r[0] for r in (await db.execute(nom_ids_subq)).all()]
            if nom_ids:
                q = q.where(and_(Document.entity_type == "nomenclature", Document.entity_id.in_(nom_ids)))
            else:
                q = q.where(Document.entity_id < 0)
        else:
            q = q.where(Document.entity_type == entity_type, Document.entity_id == eid)
    elif entity_type:
        q = q.where(Document.entity_type == entity_type)
    elif entity_id is not None:
        q = q.where(Document.entity_id == entity_id)

    if search and search.strip():
        pattern = f"%{search.strip().lower()}%"
        q = q.where(func.lower(Document.filename).like(pattern))
    q = q.order_by(Document.created_at.desc())
    result = await db.execute(q)
    items = result.scalars().all()
    return {
        "items": [
            {
                "id": d.id,
                "entity_type": d.entity_type,
                "entity_id": d.entity_id,
                "filename": d.filename,
                "mime_type": d.mime_type,
                "file_size": d.file_size,
                "created_at": d.created_at.isoformat() if d.created_at else None,
            }
            for d in items
        ]
    }


@router.get("/{doc_id}/download")
async def download_document(
    doc_id: int,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Скачать документ."""
    company_ids = await get_user_company_ids(user, db)
    q = select(Document).where(Document.id == doc_id)
    if company_ids is not None:
        q = q.where(Document.company_id.in_(company_ids))
    result = await db.execute(q)
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    root = _storage_root()
    full_path = root / doc.storage_path
    if not full_path.exists():
        raise HTTPException(404, "Файл не найден на диске")
    return FileResponse(
        path=str(full_path),
        filename=doc.filename,
        media_type=doc.mime_type,
    )


@router.get("/{doc_id}/view")
async def view_document(
    doc_id: int,
    user: User = Depends(require_trial_or_subscription),
    db: AsyncSession = Depends(get_db),
):
    """Просмотр в браузере (inline, для PDF/картинок)."""
    company_ids = await get_user_company_ids(user, db)
    q = select(Document).where(Document.id == doc_id)
    if company_ids is not None:
        q = q.where(Document.company_id.in_(company_ids))
    result = await db.execute(q)
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    root = _storage_root()
    full_path = root / doc.storage_path
    if not full_path.exists():
        raise HTTPException(404, "Файл не найден на диске")
    return FileResponse(
        path=str(full_path),
        filename=doc.filename,
        media_type=doc.mime_type,
        headers={"Content-Disposition": "inline"},
    )


@router.get("/public/{doc_id}")
async def public_view_document(
    doc_id: int,
    entity_type: str = Query(..., description="Тип сущности для проверки"),
    entity_id: int = Query(..., description="ID сущности для проверки"),
    db: AsyncSession = Depends(get_db),
):
    """Публичный просмотр документа (для страницы по QR, без авторизации)."""
    q = select(Document).where(
        Document.id == doc_id,
        Document.entity_type == entity_type,
        Document.entity_id == entity_id,
    )
    result = await db.execute(q)
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    root = _storage_root()
    full_path = root / doc.storage_path
    if not full_path.exists():
        raise HTTPException(404, "Файл не найден")
    return FileResponse(
        path=str(full_path),
        filename=doc.filename,
        media_type=doc.mime_type,
        headers={"Content-Disposition": "inline"},
    )


@router.delete("/{doc_id}")
async def delete_document(
    doc_id: int,
    user: User = Depends(require_can_delete_documents),
    db: AsyncSession = Depends(get_db),
):
    """Удалить документ."""
    company_ids = await get_user_company_ids(user, db)
    q = select(Document).where(Document.id == doc_id)
    if company_ids is not None:
        q = q.where(Document.company_id.in_(company_ids))
    result = await db.execute(q)
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    root = _storage_root()
    full_path = root / doc.storage_path
    if full_path.exists():
        try:
            full_path.unlink()
        except OSError:
            pass
    await db.delete(doc)
    await db.commit()
    return {"ok": True}
