#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Сборка Markdown из client-handover в один DOCX (python-docx + markdown + BeautifulSoup)."""

from __future__ import annotations

import sys
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

DIR = Path(__file__).resolve().parent
FILES = [
    "README.md",
    "02-functional-specification.md",
    "03-installation-and-operation.md",
    "04-lifecycle-support.md",
    "05-user-manual.md",
    "06-admin-manual.md",
    "07-software-composition.md",
]
OUT = DIR / "ikamdocs-client-handover.docx"


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_inline_runs(paragraph: Paragraph, node: Tag) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            t = str(child).strip()
            if t:
                paragraph.add_run(t)
        elif isinstance(child, Tag):
            name = child.name.lower()
            if name in ("strong", "b"):
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif name in ("em", "i"):
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            elif name == "a":
                run = paragraph.add_run(child.get_text())
                run.font.color.rgb = RGBColor(0x0, 0x66, 0xCC)
            elif name == "br":
                paragraph.add_run().add_break()
            elif name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            else:
                _add_inline_runs(paragraph, child)


def _add_paragraph_from_tag(doc: Document, tag: Tag) -> None:
    p = doc.add_paragraph()
    _add_inline_runs(p, tag)


def _add_table(doc: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr", recursive=False)
    if not rows:
        return
    cols = max(len(r.find_all(["th", "td"])) for r in rows)
    if cols == 0:
        return
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Table Grid"
    for ri, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        for ci, td in enumerate(cells):
            if ci >= cols:
                break
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            if td.find(["strong", "b"]):
                p.add_run(td.get_text()).bold = True
            else:
                p.add_run(td.get_text())
            if ri == 0 or td.name == "th":
                _set_cell_shading(cell, "E7EEF7")
    doc.add_paragraph()


def _add_block(doc: Document, tag: Tag) -> None:
    name = tag.name.lower()
    if name in ("h1", "h2", "h3", "h4"):
        level = int(name[1])
        text = tag.get_text().strip()
        if not text:
            return
        doc.add_heading(text, level=min(level, 3))
        return
    if name == "p":
        p = doc.add_paragraph()
        _add_inline_runs(p, tag)
        return
    if name == "ul":
        for li in tag.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, li)
        return
    if name == "ol":
        for li in tag.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, li)
        return
    if name == "table":
        _add_table(doc, tag)
        return
    if name == "hr":
        doc.add_paragraph("—" * 40)
        return
    if name == "blockquote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.right_indent = Pt(18)
        _add_inline_runs(p, tag)
        return
    if name == "pre":
        p = doc.add_paragraph()
        code_el = tag.find("code")
        txt = code_el.get_text() if code_el else tag.get_text()
        run = p.add_run(txt)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        return
    if name == "div":
        for child in tag.children:
            if isinstance(child, NavigableString):
                continue
            if isinstance(child, Tag):
                _add_block(doc, child)
        return


def md_to_html(md_text: str) -> str:
    return markdown.markdown(
        md_text,
        extensions=[
            "tables",
            "fenced_code",
            "nl2br",
            "sane_lists",
        ],
    )


def html_to_docx_fragment(doc: Document, html: str) -> None:
    soup = BeautifulSoup(f"<div>{html}</div>", "lxml")
    root = soup.find("div")
    if not root:
        return
    for child in root.children:
        if isinstance(child, NavigableString):
            continue
        if isinstance(child, Tag):
            _add_block(doc, child)


def strip_front_matter(text: str) -> str:
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            return parts[2].lstrip("\n")
    return text


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("ikamdocs — документация для передачи клиенту", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.add_run("Разделы 2–7 по запросу + оглавление (README)").italic = True
    doc.add_paragraph()

    for fname in FILES:
        path = DIR / fname
        if not path.exists():
            print(f"Skip missing: {path}", file=sys.stderr)
            continue
        raw = path.read_text(encoding="utf-8")
        raw = strip_front_matter(raw)
        html = md_to_html(raw)
        html_to_docx_fragment(doc, html)
        doc.add_page_break()

    # remove trailing page break
    if doc.element.body[-1].tag.endswith("sectPr"):
        pass
    # last element might be last paragraph break - acceptable

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    build()
